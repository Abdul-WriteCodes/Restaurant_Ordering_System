"""
Exploratory Factor Analysis and Comfirmatory Factor Analysis

"""

import io
import zipfile
import warnings
from datetime import datetime

import numpy as np
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import streamlit as st
from factor_analyzer import FactorAnalyzer, calculate_kmo, calculate_bartlett_sphericity
import gspread
from google.oauth2.service_account import Credentials

# ── sklearn >= 1.6 compatibility ─────────────────────────────────────────────
try:
    import factor_analyzer.factor_analyzer as _fa_mod
    import factor_analyzer.confirmatory_factor_analyzer as _cfa_mod
    from sklearn.utils.validation import check_array as _orig_check_array
    def _compat_check_array(*args, **kwargs):
        if "force_all_finite" in kwargs:
            val = kwargs.pop("force_all_finite")
            kwargs["ensure_all_finite"] = (val is True)
        return _orig_check_array(*args, **kwargs)
    _fa_mod.check_array  = _compat_check_array
    _cfa_mod.check_array = _compat_check_array
except Exception:
    pass
# ─────────────────────────────────────────────────────────────────────────────

warnings.filterwarnings("ignore")

# ══════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="EFActor — Psychometric Analysis Platform",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════
# THEME CONSTANTS
# ══════════════════════════════════════════════════════════════════
C = dict(
    accent="#6c8dfa", accent2="#a78bfa",
    green="#34d399", red="#f87171", yellow="#fbbf24",
    bg="#0a0c14", surface="#111520", surface2="#181d2e", border="#1e2540",
    text="#e2e8f0", muted="#64748b",
)
LAYOUT_BASE = dict(
    paper_bgcolor=C["surface"], plot_bgcolor=C["surface"],
    font=dict(color=C["text"], family="Inter, sans-serif"),
    margin=dict(l=40, r=20, t=50, b=40),
    colorway=[C["accent"], C["accent2"], C["green"], C["yellow"], C["red"]],
)

# ══════════════════════════════════════════════════════════════════
# GLOBAL CSS
# ══════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap');
:root {
  --accent:#6c8dfa; --accent2:#a78bfa; --green:#34d399;
  --red:#f87171; --yellow:#fbbf24; --border:#1e2540;
  --surface:#111520; --surface2:#181d2e; --muted:#64748b;
  --bg:#0a0c14; --text:#e2e8f0;
}
html,body,[class*="css"]{ font-family:'Inter',system-ui,sans-serif; background-color:var(--bg); }
.stApp{ background:var(--bg); }
[data-testid="stSidebar"]{ background:#0d1020 !important; border-right:1px solid var(--border); }
h1{ color:var(--accent) !important; font-weight:800 !important; letter-spacing:-0.5px; }
h2{ color:var(--accent2) !important; border-bottom:1px solid var(--border); padding-bottom:8px; font-weight:700 !important; }
h3{ color:var(--text) !important; font-weight:600 !important; }
.step-badge{ display:inline-flex; align-items:center; gap:10px;
  background:linear-gradient(135deg,var(--surface),var(--surface2));
  border:1px solid var(--border); border-left:3px solid var(--accent);
  padding:11px 18px; border-radius:8px; color:var(--text);
  font-size:.95rem; font-weight:600; margin-bottom:18px; width:100%; }
.step-num{ background:linear-gradient(135deg,var(--accent),var(--accent2));
  color:var(--bg); border-radius:50%; width:26px; height:26px;
  display:inline-flex; align-items:center; justify-content:center;
  font-size:.8rem; font-weight:700; flex-shrink:0; }
.pill-pass{ background:rgba(52,211,153,.12); color:var(--green); border:1px solid rgba(52,211,153,.4);
  border-radius:20px; padding:3px 14px; font-size:.78rem; font-weight:600; display:inline-block; }
.pill-fail{ background:rgba(248,113,113,.12); color:var(--red); border:1px solid rgba(248,113,113,.4);
  border-radius:20px; padding:3px 14px; font-size:.78rem; font-weight:600; display:inline-block; }
.metric-card{ background:var(--surface); border:1px solid var(--border);
  border-radius:10px; padding:18px 20px; text-align:center; position:relative; overflow:hidden; }
.metric-card::before{ content:''; position:absolute; top:0; left:0; right:0; height:2px;
  background:linear-gradient(90deg,var(--accent),var(--accent2)); }
.metric-val{ font-size:2rem; font-weight:700; color:var(--accent); line-height:1.1; }
.metric-label{ font-size:.7rem; color:var(--muted); text-transform:uppercase;
  letter-spacing:1px; margin-top:6px; font-weight:500; }
.info-box{ background:rgba(108,141,250,.07); border:1px solid rgba(108,141,250,.2);
  border-radius:8px; padding:12px 16px; color:#c7d2fe; font-size:.88rem; margin:8px 0; line-height:1.6; }
.warn-box{ background:rgba(251,191,36,.07); border:1px solid rgba(251,191,36,.25);
  border-radius:8px; padding:12px 16px; color:#fde68a; font-size:.88rem; margin:8px 0; line-height:1.6; }
.lock-box{ background:rgba(108,141,250,.06); border:1px solid rgba(108,141,250,.2);
  border-left:3px solid var(--accent); border-radius:8px; padding:18px 20px; margin:8px 0; }
.section-divider{ border:none; border-top:1px solid var(--border); margin:32px 0; }
code,pre{ font-family:'JetBrains Mono',monospace !important; }
.stButton>button{ background:linear-gradient(135deg,#151a2e,#1e2745) !important;
  border:1px solid var(--accent) !important; color:var(--accent) !important;
  font-family:'Inter',sans-serif !important; font-size:.92rem !important;
  font-weight:600 !important; border-radius:8px !important; }
.stButton>button:hover{ background:linear-gradient(135deg,var(--accent),#4f6fe8) !important;
  color:var(--bg) !important; border-color:transparent !important; }
.stDownloadButton>button{ background:linear-gradient(135deg,#0f1f18,#122518) !important;
  border:1px solid var(--green) !important; color:var(--green) !important;
  border-radius:8px !important; font-weight:600 !important; }
.stTabs [data-baseweb="tab-list"]{ background:var(--surface); border-radius:10px;
  padding:4px; gap:4px; border:1px solid var(--border); }
.stTabs [data-baseweb="tab"]{ background:transparent; border-radius:8px; color:var(--muted);
  font-size:13px; font-weight:500; padding:8px 18px; border:none; }
.stTabs [aria-selected="true"]{ background:linear-gradient(135deg,var(--accent),#4f6fe8) !important; color:white !important; }
.stTextInput input,.stNumberInput input{ background:var(--surface2) !important;
  border:1px solid var(--border) !important; border-radius:8px !important; color:var(--text) !important; }
#MainMenu{ visibility:hidden; } footer{ visibility:hidden; }
header[data-testid="stHeader"]{ background:transparent !important; }
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# GOOGLE SHEETS / CREDIT ENGINE
# ══════════════════════════════════════════════════════════════════
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]
SHEET_NAME       = "Sheet1"
COL_KEY          = "Key"
COL_CREDITS      = "Credits"
COL_OWNER        = "Email"
COL_ISSUED       = "DatePurchased"
REQUIRED_HEADERS = [COL_KEY, COL_CREDITS, COL_ISSUED, COL_OWNER]


@st.cache_resource(show_spinner=False)
def _get_gsheet_client():
    creds_dict = dict(st.secrets["gcp_service_account"])
    creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
    return gspread.authorize(creds)


def _get_keys_worksheet():
    gc = _get_gsheet_client()
    sheet_id = st.secrets["EFACTOR_SHEET_ID"]
    sh = gc.open_by_key(sheet_id)
    return sh.worksheet(SHEET_NAME)


def validate_key(access_key: str):
    try:
        ws = _get_keys_worksheet()
        records = ws.get_all_records(
            expected_headers=REQUIRED_HEADERS,
            value_render_option="UNFORMATTED_VALUE",
        )
        records = [r for r in records if any(str(v).strip() for v in r.values())]
        for row in records:
            if str(row.get(COL_KEY, "")).strip() == access_key.strip():
                return row
        return None
    except Exception as e:
        st.error(f"Key validation error: {e}")
        return None


def get_credits(access_key: str) -> int:
    try:
        ws = _get_keys_worksheet()
        records = ws.get_all_records(
            expected_headers=REQUIRED_HEADERS,
            value_render_option="UNFORMATTED_VALUE",
        )
        records = [r for r in records if any(str(v).strip() for v in r.values())]
        for row in records:
            if str(row.get(COL_KEY, "")).strip() == access_key.strip():
                return int(row.get(COL_CREDITS, 0))
        return 0
    except Exception:
        return 0


def deduct_credits(access_key: str, amount: int) -> int:
    try:
        ws = _get_keys_worksheet()
        records = ws.get_all_records(
            expected_headers=REQUIRED_HEADERS,
            value_render_option="UNFORMATTED_VALUE",
        )
        records = [r for r in records if any(str(v).strip() for v in r.values())]
        header = ws.row_values(1)
        credits_col_idx = header.index(COL_CREDITS) + 1
        for i, row in enumerate(records):
            if str(row.get(COL_KEY, "")).strip() == access_key.strip():
                row_number = i + 2
                current  = int(row.get(COL_CREDITS, 0))
                new_val  = max(0, current - amount)
                ws.update_cell(row_number, credits_col_idx, new_val)
                return new_val
        return 0
    except Exception as e:
        st.error(f"Credit deduction error: {e}")
        return 0


def export_credit_cost(n_rows: int) -> int:
    if n_rows <= 300:
        return 1
    elif n_rows <= 1000:
        return 2
    else:
        return 5


# ══════════════════════════════════════════════════════════════════
# PLANS CONFIG
# ══════════════════════════════════════════════════════════════════
PLANS = [
    {
        "name": "Starter", "price": "$9", "credits": "10 credits",
        "period": "one-time", "color": "#34d399", "highlight": False,
        "features": ["Unlimited EFA & CFA runs","Dataset exports",
                     "Synthetic data generation",
                     "Word analysis report (.docx)","Access key via email"],
        "link": "https://flutterwave.com/pay/blhpfnlv6qj0",
    },
    {
        "name": "Standard", "price": "$19", "credits": "30 credits",
        "period": "one-time", "color": "#6c8dfa", "highlight": True,
        "features": ["Unlimited EFA & CFA runs","Dataset exports",
                     "Synthetic data generation",
                     "Word analysis report (.docx)","Access key via email"],
        "link": "https://flutterwave.com/pay/rdzxbnicgu4f",
    },
    {
        "name": "Pro", "price": "$49", "credits": "100 credits",
        "period": "one-time", "color": "#a78bfa", "highlight": False,
        "features": ["Unlimited EFA & CFA runs","Dtaset exports",
                     "Unlimited rows per export","Synthetic data generation",
                     "Word analysis report (.docx)",
                     "Access key via email"],
        "link": "https://flutterwave.com/pay/skqi0vxinkrx",
    },
]


# ══════════════════════════════════════════════════════════════════
# LANDING PAGE
# ══════════════════════════════════════════════════════════════════
def render_landing():
    st.markdown("""
    <style>
    section[data-testid="stSidebar"]{ display:none !important; }
    .block-container{ max-width:1020px; margin:0 auto; padding-top:0rem; }
    </style>""", unsafe_allow_html=True)

    # Hero
    st.markdown("""
    <div style="text-align:center;padding:72px 20px 48px;">
      <div style="display:inline-flex;align-items:center;gap:14px;margin-bottom:32px;">
        <div style="width:44px;height:44px;background:linear-gradient(135deg,#6c8dfa,#a78bfa);
                    border-radius:10px;display:flex;align-items:center;justify-content:center;
                    font-size:22px;flex-shrink:0;">🔬</div>
        <span style="font-family:Inter,sans-serif;font-size:26px;font-weight:800;
                     background:linear-gradient(90deg,#6c8dfa,#a78bfa);
                     -webkit-background-clip:text;-webkit-text-fill-color:transparent;
                     letter-spacing:-0.5px;">EFActor</span>
      </div>
      <div style="font-family:Inter,sans-serif;font-size:11px;color:#34d399;
                  letter-spacing:3px;text-transform:uppercase;margin-bottom:20px;font-weight:500;">
        Psychometric Analysis · EFA · CFA · Synthetic Data
      </div>
      <h1 style="font-family:Inter,sans-serif;font-size:clamp(28px,4vw,52px);font-weight:800;
                 background:linear-gradient(135deg,#e2e8f0 40%,#a78bfa);
                 -webkit-background-clip:text;-webkit-text-fill-color:transparent;
                 line-height:1.1;margin:0 0 24px;letter-spacing:-1px;">
        Factor Analysis,<br>Built for Researchers
      </h1>
      <p style="font-family:Inter,sans-serif;font-size:17px;color:#94a3b8;
                max-width:560px;margin:0 auto 52px;line-height:1.8;">
        Run EFA and CFA on your survey data. Diagnose item issues, confirm
        factor structure, generate synthetic datasets — all in one rigorous pipeline.
      </p>
    </div>
    """, unsafe_allow_html=True)

    # How it works
    st.markdown("""
    <div style="display:grid;grid-template-columns:repeat(3,1fr);gap:16px;
                max-width:860px;margin:0 auto 72px;">
      <div style="background:#111520;border:1px solid #1e2540;border-radius:14px;padding:28px 22px;text-align:center;">
        <div style="font-size:28px;margin-bottom:14px;">⬆</div>
        <div style="font-family:Inter,sans-serif;font-size:14px;font-weight:700;color:#e2e8f0;margin-bottom:8px;">Upload</div>
        <div style="font-family:Inter,sans-serif;font-size:13px;color:#64748b;line-height:1.7;">
          Upload your CSV or Excel survey dataset. Numeric columns detected automatically.
        </div>
      </div>
      <div style="background:#111520;border:1px solid #1e2540;border-top:2px solid #6c8dfa;
                  border-radius:14px;padding:28px 22px;text-align:center;">
        <div style="font-size:28px;margin-bottom:14px;">◈</div>
        <div style="font-family:Inter,sans-serif;font-size:14px;font-weight:700;color:#e2e8f0;margin-bottom:8px;">Analyse</div>
        <div style="font-family:Inter,sans-serif;font-size:13px;color:#64748b;line-height:1.7;">
          Run KMO, Bartlett, EFA and CFA with full diagnostics. Iterate freely — no limits.
        </div>
      </div>
      <div style="background:#111520;border:1px solid #1e2540;border-radius:14px;padding:28px 22px;text-align:center;">
        <div style="font-size:28px;margin-bottom:14px;">⬇</div>
        <div style="font-family:Inter,sans-serif;font-size:14px;font-weight:700;color:#e2e8f0;margin-bottom:8px;">Export</div>
        <div style="font-family:Inter,sans-serif;font-size:13px;color:#64748b;line-height:1.7;">
          Download cleaned data, synthetic datasets and a full Word analysis report (.docx).
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Free Trial CTA
    st.markdown("""
    <div style="max-width:860px;margin:0 auto 16px;padding:0 4px;">
      <div style="background:linear-gradient(135deg,rgba(108,141,250,0.07),rgba(167,139,250,0.06));
                  border:1px solid rgba(108,141,250,0.25);border-radius:18px;
                  padding:40px 36px;position:relative;overflow:hidden;text-align:center;">
        <div style="position:absolute;top:0;left:0;right:0;height:2px;
                    background:linear-gradient(90deg,transparent,#6c8dfa 40%,#a78bfa 60%,transparent);"></div>
        <div style="font-family:Inter,sans-serif;font-size:10px;color:#34d399;
                    letter-spacing:3px;text-transform:uppercase;margin-bottom:16px;font-weight:600;">
          🔬 No credit card · Instant access
        </div>
        <div style="font-family:Inter,sans-serif;font-size:26px;font-weight:800;
                    background:linear-gradient(90deg,#e2e8f0,#6c8dfa);
                    -webkit-background-clip:text;-webkit-text-fill-color:transparent;
                    margin-bottom:12px;line-height:1.2;">Try EFActor Free</div>
        <p style="font-family:Inter,sans-serif;font-size:13px;color:#94a3b8;
                  max-width:500px;margin:0 auto 28px;line-height:1.9;">
          Upload data, run full EFA and CFA, generate synthetic datasets, preview all results.
          Export is unlocked with a paid access key.
        </p>
        <div style="display:flex;flex-wrap:wrap;justify-content:center;gap:10px;">
          <span style="background:rgba(52,211,153,.1);border:1px solid rgba(52,211,153,.25);color:#34d399;
                       font-family:Inter,sans-serif;font-size:11px;font-weight:500;padding:6px 14px;border-radius:8px;">✓ Upload CSV / Excel</span>
          <span style="background:rgba(52,211,153,.1);border:1px solid rgba(52,211,153,.25);color:#34d399;
                       font-family:Inter,sans-serif;font-size:11px;font-weight:500;padding:6px 14px;border-radius:8px;">✓ Unlimited EFA &amp; CFA runs</span>
          <span style="background:rgba(52,211,153,.1);border:1px solid rgba(52,211,153,.25);color:#34d399;
                       font-family:Inter,sans-serif;font-size:11px;font-weight:500;padding:6px 14px;border-radius:8px;">✓ Full diagnostics &amp; plots</span>
          <span style="background:rgba(52,211,153,.1);border:1px solid rgba(52,211,153,.25);color:#34d399;
                       font-family:Inter,sans-serif;font-size:11px;font-weight:500;padding:6px 14px;border-radius:8px;">✓ Synthetic data preview</span>
          <span style="background:rgba(248,113,113,.07);border:1px solid rgba(248,113,113,.2);color:#f87171;
                       font-family:Inter,sans-serif;font-size:11px;font-weight:500;padding:6px 14px;border-radius:8px;opacity:.85;">✗ Data export (requires credits)</span>
          <span style="background:rgba(248,113,113,.07);border:1px solid rgba(248,113,113,.2);color:#f87171;
                       font-family:Inter,sans-serif;font-size:11px;font-weight:500;padding:6px 14px;border-radius:8px;opacity:.85;">✗ Word report (.docx) requires credits</span>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    _, trial_col, _ = st.columns([1.4, 1, 1.4])
    with trial_col:
        trial_btn = st.button("◈ Start Free Trial", type="primary", width='stretch', key="trial_btn")
    st.markdown("<div style='margin-bottom:56px;'></div>", unsafe_allow_html=True)

    if trial_btn:
        st.session_state.update({
            "authenticated": True, "is_free_trial": True,
            "access_key": "FREE-TRIAL", "key_owner": "Free Trial", "credits": 0,
        })
        st.rerun()

    # Pricing header
    st.markdown("""
    <div style="text-align:center;margin-bottom:40px;">
      <div style="font-family:Inter,sans-serif;font-size:10px;color:#64748b;
                  letter-spacing:3px;text-transform:uppercase;margin-bottom:12px;font-weight:500;">Pricing</div>
      <div style="font-family:Inter,sans-serif;font-size:30px;font-weight:800;color:#e2e8f0;letter-spacing:-0.5px;">
        Simple, credit-based access
      </div>
      <div style="font-family:Inter,sans-serif;font-size:14px;color:#64748b;margin-top:10px;line-height:1.7;">
        Iterate freely. Credits are only used when you export.<br>
        Access key delivered to your email after purchase.
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Pricing cards
    p_cols = st.columns(3, gap="medium")
    for col, plan in zip(p_cols, PLANS):
        border_top = f"border-top:3px solid {plan['color']};" if plan["highlight"] else f"border-top:2px solid {plan['color']}44;"
        shadow = f"box-shadow:0 16px 48px {plan['color']}20;" if plan["highlight"] else ""
        badge_html = (
            '<div style="background:rgba(108,141,250,0.12);color:#6c8dfa;font-family:Inter,sans-serif;'
            'font-size:9px;font-weight:700;letter-spacing:2px;text-transform:uppercase;padding:4px 12px;'
            'border-radius:20px;display:inline-block;margin-bottom:16px;border:1px solid rgba(108,141,250,.3);">★ MOST POPULAR</div>'
            if plan["highlight"] else "<div style='height:28px;'></div>"
        )
        feats = "".join(
            f'<div style="font-family:Inter,sans-serif;font-size:12px;color:#94a3b8;'
            f'padding:7px 0;border-bottom:1px solid #1a1f35;">'
            f'<span style="color:{plan["color"]};margin-right:8px;font-weight:600;">✓</span>{f}</div>'
            for f in plan["features"]
        )
        with col:
            st.markdown(f"""
            <div style="background:#0d1020;border:1px solid #1e2540;{border_top}
                        border-radius:16px;padding:32px 26px 16px;{shadow}
                        min-height:420px;box-sizing:border-box;">
              {badge_html}
              <div style="font-family:Inter,sans-serif;font-size:17px;font-weight:700;color:#e2e8f0;margin-bottom:6px;">{plan["name"]}</div>
              <div style="font-family:Inter,sans-serif;font-size:46px;font-weight:800;color:{plan["color"]};line-height:1;margin:10px 0 4px;">{plan["price"]}</div>
              <div style="font-family:Inter,sans-serif;font-size:11px;color:#64748b;margin-bottom:24px;font-weight:500;">{plan["credits"]} · {plan["period"]}</div>
              <div>{feats}</div>
            </div>
            """, unsafe_allow_html=True)
            st.link_button(f"Get {plan['name']} →", url=plan["link"], width='stretch')

    # Credit cost table
    st.markdown("""
    <div style="max-width:860px;margin:56px auto 16px;">
      <div style="text-align:center;margin-bottom:24px;">
        <div style="font-family:Inter,sans-serif;font-size:16px;font-weight:700;color:#e2e8f0;margin-bottom:4px;">Export Credit Costs</div>
        <div style="font-family:Inter,sans-serif;font-size:12px;color:#64748b;">Analysis runs are always free. Credits only deducted on export.</div>
      </div>
      <div style="display:grid;grid-template-columns:repeat(4,1fr);gap:12px;">
        <div style="background:#111520;border:1px solid #1e2540;border-radius:12px;padding:20px 16px;text-align:center;">
          <div style="font-size:22px;font-weight:800;color:#34d399;margin-bottom:4px;">1</div>
          <div style="font-size:10px;color:#64748b;font-weight:500;letter-spacing:.5px;text-transform:uppercase;margin-bottom:6px;">Credit</div>
          <div style="font-size:12px;color:#94a3b8;line-height:1.6;">Data export<br>≤ 300 rows</div>
        </div>
        <div style="background:#111520;border:1px solid #1e2540;border-radius:12px;padding:20px 16px;text-align:center;">
          <div style="font-size:22px;font-weight:800;color:#6c8dfa;margin-bottom:4px;">2</div>
          <div style="font-size:10px;color:#64748b;font-weight:500;letter-spacing:.5px;text-transform:uppercase;margin-bottom:6px;">Credits</div>
          <div style="font-size:12px;color:#94a3b8;line-height:1.6;">Data export<br>301 – 1,000 rows</div>
        </div>
        <div style="background:#111520;border:1px solid #1e2540;border-radius:12px;padding:20px 16px;text-align:center;">
          <div style="font-size:22px;font-weight:800;color:#a78bfa;margin-bottom:4px;">5</div>
          <div style="font-size:10px;color:#64748b;font-weight:500;letter-spacing:.5px;text-transform:uppercase;margin-bottom:6px;">Credits</div>
          <div style="font-size:12px;color:#94a3b8;line-height:1.6;">Data export<br>&gt; 1,000 rows</div>
        </div>
        <div style="background:#111520;border:1px solid #1e2540;border-radius:12px;padding:20px 16px;text-align:center;">
          <div style="font-size:22px;font-weight:800;color:#fbbf24;margin-bottom:4px;">1</div>
          <div style="font-size:10px;color:#64748b;font-weight:500;letter-spacing:.5px;text-transform:uppercase;margin-bottom:6px;">Credit</div>
          <div style="font-size:12px;color:#94a3b8;line-height:1.6;">HTML analysis<br>report</div>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Access key entry
    st.markdown("""
    <div style="text-align:center;margin-top:72px;margin-bottom:24px;">
      <div style="font-family:Inter,sans-serif;font-size:13px;font-weight:600;
                  color:#64748b;letter-spacing:1px;text-transform:uppercase;margin-bottom:6px;">
        Already have an access key?
      </div>
      <div style="font-family:Inter,sans-serif;font-size:12px;color:#374151;">
        Your key is emailed to you after purchase.
      </div>
    </div>
    """, unsafe_allow_html=True)

    _, c, _ = st.columns([1.2, 1.6, 1.2])
    with c:
        key_input = st.text_input("Access Key", placeholder="EFA-XXXX-XXXX-XXXX",
                                   label_visibility="collapsed", key="landing_key_input")
        enter_btn = st.button("🔬 Enter EFActor", type="primary", width='stretch')
        if enter_btn:
            if not key_input.strip():
                st.error("Please enter your access key.")
            else:
                with st.spinner("Validating key…"):
                    row = validate_key(key_input.strip())
                if row is None:
                    st.error("Invalid access key. Please check and try again.")
                elif int(row.get(COL_CREDITS, 0)) <= 0:
                    st.error("This key has 0 credits remaining. Please purchase a new plan.")
                else:
                    st.session_state.update({
                        "authenticated": True, "is_free_trial": False,
                        "access_key": key_input.strip(),
                        "key_owner": row.get(COL_OWNER, "Researcher"),
                        "credits": int(row.get(COL_CREDITS, 0)),
                    })
                    st.rerun()

    # Footer
    st.markdown("""
    <div style="text-align:center;margin-top:40px;padding-bottom:16px;">
      <a href="https://x.com/Bayantx360" target="_blank"
         style="font-family:Inter,sans-serif;font-size:12px;font-weight:600;color:white;
                background:linear-gradient(135deg,#151a2e,#1e2745);
                border:1px solid #1e2540;padding:10px 18px;border-radius:8px;
                margin:0 6px;display:inline-block;text-decoration:none;">👤 Get Access Key</a>
      <a href="mailto:Bayantx360@gmail.com"
         style="font-family:Inter,sans-serif;font-size:12px;font-weight:600;color:white;
                background:linear-gradient(135deg,#151a2e,#1e2745);
                border:1px solid #1e2540;padding:10px 18px;border-radius:8px;
                margin:0 6px;display:inline-block;text-decoration:none;">⚙️ Support</a>
    </div>
    <div style="text-align:center;font-family:Inter,sans-serif;font-size:10px;
                color:#1e2540;margin-top:40px;letter-spacing:1px;padding-bottom:40px;">
      © 2025 EFActor · Psychometric Analysis Platform • Bayantx360
    </div>
    """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# AUTH GATE
# ══════════════════════════════════════════════════════════════════
if not st.session_state.get("authenticated", False):
    render_landing()
    st.stop()

# Refresh live credit balance (paid users only)
if not st.session_state.get("is_free_trial", False):
    st.session_state["credits"] = get_credits(st.session_state["access_key"])


# ══════════════════════════════════════════════════════════════════
# EFA FUNCTIONS  (core logic unchanged)
# ══════════════════════════════════════════════════════════════════

def _ensure_psd_dataframe(df: pd.DataFrame, eps: float = 1e-4) -> pd.DataFrame:
    """
    Return a copy of df whose correlation matrix is positive-definite.
    Strategy:
      1. Drop any constant columns (std == 0) — they break corr().
      2. Regularise the correlation matrix via eigenvalue flooring.
      3. Reconstruct df columns so they honour the regularised corr matrix
         while preserving original means and standard deviations.
    This is called as a safety net before any FactorAnalyzer.fit() call.
    """
    df = df.copy()
    # Drop zero-variance columns (can't be rescued here — caller should handle)
    df = df.loc[:, df.std() > 1e-9]
    if df.shape[1] < 2:
        return df

    # Check if already PD
    corr = df.corr().values
    eigvals = np.linalg.eigvalsh(corr)
    if eigvals.min() > eps:
        return df   # already fine

    # Floor negative / near-zero eigenvalues
    ev, evec = np.linalg.eigh(corr)
    ev_fixed = np.maximum(ev, eps)
    corr_fixed = evec @ np.diag(ev_fixed) @ evec.T
    # Re-normalise to valid correlation matrix (diagonal = 1)
    d = np.sqrt(np.diag(corr_fixed))
    corr_fixed = corr_fixed / np.outer(d, d)
    np.fill_diagonal(corr_fixed, 1.0)

    # Reconstruct data from Cholesky of fixed corr × original stds/means
    try:
        L = np.linalg.cholesky(corr_fixed + np.eye(len(corr_fixed)) * eps)
    except np.linalg.LinAlgError:
        # Last resort: add diagonal ridge until it works
        ridge = eps
        for _ in range(20):
            try:
                L = np.linalg.cholesky(corr_fixed + np.eye(len(corr_fixed)) * ridge)
                break
            except np.linalg.LinAlgError:
                ridge *= 10
        else:
            return df   # give up, return as-is

    rng = np.random.default_rng(0)
    z = rng.standard_normal((len(df), len(df.columns)))
    z_corr = z @ L.T
    # Rescale each column to original mean/std
    stds  = df.std().values
    means = df.mean().values
    z_norm = (z_corr - z_corr.mean(axis=0)) / np.maximum(z_corr.std(axis=0), 1e-9)
    reconstructed = z_norm * stds + means
    return pd.DataFrame(reconstructed, columns=df.columns, index=df.index)


def check_efa_suitability(df):
    kmo_all, kmo_model = calculate_kmo(df)
    chi2, p = calculate_bartlett_sphericity(df)
    labels = {.9:"Marvellous",.8:"Meritorious",.7:"Middling",.6:"Mediocre",.5:"Miserable"}
    kmo_label = next((v for k, v in sorted(labels.items(), reverse=True) if kmo_model >= k), "Unacceptable")
    return dict(kmo_model=round(float(kmo_model),4), kmo_label=kmo_label, kmo_pass=kmo_model>=0.6,
                bartlett_chi2=round(float(chi2),4), bartlett_p=round(float(p),6),
                bartlett_pass=p<0.05, overall_pass=kmo_model>=0.6 and p<0.05)

def determine_n_factors(df):
    df = _ensure_psd_dataframe(df)
    max_factors = max(1, min(len(df.columns)-1, len(df)-1))
    try:
        fa = FactorAnalyzer(n_factors=max_factors, rotation=None)
        fa.fit(df)
        ev, _ = fa.get_eigenvalues()
    except Exception:
        # Fallback: use eigenvalues of correlation matrix directly
        corr = df.corr().values
        ev = np.sort(np.linalg.eigvalsh(corr))[::-1]
    return dict(eigenvalues=ev.tolist(), suggested_n=max(1, int(np.sum(np.array(ev) > 1))))

def run_efa(df, n_factors, rotation="varimax"):
    df = _ensure_psd_dataframe(df)
    n_factors = min(n_factors, len(df.columns)-1)
    try:
        fa = FactorAnalyzer(n_factors=n_factors, rotation=rotation)
        fa.fit(df)
        factor_labels = [f"F{i+1}" for i in range(n_factors)]
        loadings = pd.DataFrame(fa.loadings_, index=df.columns, columns=factor_labels)
        communalities = pd.Series(fa.get_communalities(), index=df.columns, name="Communality")
        variance = pd.DataFrame(fa.get_factor_variance(),
                                index=["SS Loadings","Proportion Var","Cumulative Var"],
                                columns=factor_labels).T
    except Exception:
        # Fallback: PCA-based loadings from eigendecomposition of corr matrix
        corr = df.corr().values
        ev, evec = np.linalg.eigh(corr)
        idx = np.argsort(ev)[::-1][:n_factors]
        ev_top = np.maximum(ev[idx], 0)
        raw_loadings = evec[:, idx] * np.sqrt(ev_top)
        factor_labels = [f"F{i+1}" for i in range(n_factors)]
        loadings = pd.DataFrame(raw_loadings, index=df.columns, columns=factor_labels)
        communalities = pd.Series(
            np.clip((raw_loadings ** 2).sum(axis=1), 0, 1),
            index=df.columns, name="Communality"
        )
        variance = pd.DataFrame(
            np.zeros((n_factors, 3)),
            index=factor_labels,
            columns=["SS Loadings","Proportion Var","Cumulative Var"]
        )
    return dict(loadings=loadings, communalities=communalities, variance=variance, n_factors=n_factors)

def diagnose_loadings(loadings, communalities, load_thresh=0.4, comm_thresh=0.3):
    records = []
    for var in loadings.index:
        abs_row = np.abs(loadings.loc[var])
        max_load = abs_row.max()
        n_high = int((abs_row >= load_thresh).sum())
        comm = float(communalities[var])
        issues, severity = [], 0.0
        if comm < comm_thresh:
            issues.append("Low Communality"); severity += (comm_thresh - comm) * 3
        if n_high == 0:
            issues.append("Weak Loader"); severity += (load_thresh - max_load) * 2
        elif n_high > 1:
            issues.append("Cross-Loader")
            s = sorted(abs_row.values, reverse=True)
            severity += (1 - (s[0] - s[1])) * 2
        records.append(dict(Variable=var, MaxLoading=round(max_load,4),
                            FactorsAboveThreshold=n_high, Communality=round(comm,4),
                            Issue=", ".join(issues) if issues else "OK",
                            Severity=round(severity,4), RecommendDrop=len(issues)>0))
    return pd.DataFrame(records).sort_values("Severity", ascending=False).reset_index(drop=True)


# ══════════════════════════════════════════════════════════════════
# EFA AUTO-FIX ENGINE  (iterative, EFA-aware)
# ══════════════════════════════════════════════════════════════════

def _winsorize(series: pd.Series, limits=(0.05, 0.05)) -> pd.Series:
    lo = series.quantile(limits[0])
    hi = series.quantile(1.0 - limits[1])
    return series.clip(lower=lo, upper=hi)

def _log_transform(series: pd.Series) -> pd.Series:
    shifted = series - series.min() + 1e-6
    return np.log1p(shifted)

def _sqrt_transform(series: pd.Series) -> pd.Series:
    shifted = series - series.min()
    return np.sqrt(shifted)

def _add_jitter(series: pd.Series, seed=42, scale=0.01) -> pd.Series:
    rng = np.random.default_rng(seed)
    std = series.std()
    noise_scale = max(std * scale, 1e-6)
    return series + rng.normal(0, noise_scale, size=len(series))

def _rescale_to_original(fixed: pd.Series, original: pd.Series) -> pd.Series:
    """Restore original mean and SD after any transformation."""
    orig_mean, orig_std = float(original.mean()), float(original.std())
    fixed_std = float(fixed.std())
    if fixed_std > 1e-9 and orig_std > 1e-9:
        fixed = (fixed - fixed.mean()) / fixed_std * orig_std + orig_mean
    return fixed

def _detect_data_issues(df: pd.DataFrame) -> dict:
    """
    Return a dict {var: [issue_str, ...]} for every variable with detectable
    data-level problems that can cause EFA failure.
    Checks: zero-variance, outliers, skewness, kurtosis, near-perfect collinearity.
    """
    issues = {}
    corr = df.corr().abs()
    for var in df.columns:
        s = df[var]
        vi = []
        if s.std() < 1e-6:
            vi.append("zero_variance")
            issues[var] = vi
            continue
        z = (s - s.mean()) / s.std()
        n_out = int((z.abs() > 3).sum())
        if n_out > 0:
            vi.append(f"outliers:{n_out}")
        skw = float(s.skew())
        if abs(skw) > 2:
            vi.append(f"skewness:{skw:.3f}")
        krt = float(s.kurt())
        if krt > 7:
            vi.append(f"kurtosis:{krt:.3f}")
        others = corr[var].drop(var)
        if len(others) and others.max() >= 0.95:
            partner = others.idxmax()
            vi.append(f"collinear:{partner}:{others.max():.3f}")
        if vi:
            issues[var] = vi
    return issues

def _apply_fixes_for_issues(series: pd.Series, issue_tags: list,
                             original: pd.Series, seed: int, iteration: int) -> tuple:
    """
    Apply ALL relevant fixes for the given issue tags in a single pass.
    Returns (fixed_series, [fix_description, ...]).
    On later iterations, escalate aggression (tighter winsorize, larger jitter).
    """
    fixed = series.copy()
    applied = []
    # Escalate on later iterations
    winsor_pct  = max(0.01, 0.05 - (iteration - 1) * 0.01)   # 5%→4%→3%→2%→1%
    jitter_scale= 0.01 * (1 + (iteration - 1) * 0.5)          # 1%→1.5%→2%→2.5%

    has_zero    = any("zero_variance" in t for t in issue_tags)
    has_outlier = any("outliers"      in t for t in issue_tags)
    has_skew    = any("skewness"      in t for t in issue_tags)
    has_kurt    = any("kurtosis"      in t for t in issue_tags)
    has_coll    = any("collinear"     in t for t in issue_tags)

    if has_zero:
        fixed = _add_jitter(fixed, seed=seed, scale=jitter_scale)
        applied.append(f"Jitter (zero-variance, iter {iteration})")

    if has_outlier:
        fixed = _winsorize(fixed, limits=(winsor_pct, winsor_pct))
        applied.append(f"Winsorize {winsor_pct*100:.0f}th–{(1-winsor_pct)*100:.0f}th pctile (iter {iteration})")

    if has_skew:
        skw = float(fixed.skew())
        if skw > 2:
            fixed = _log_transform(fixed)
            applied.append(f"Log1p transform (pos skew {skw:.2f}, iter {iteration})")
        elif skw < -2:
            fixed = _log_transform(fixed.max() - fixed)
            applied.append(f"Reflected log (neg skew {skw:.2f}, iter {iteration})")

    if has_kurt and not has_outlier:
        kurt_winsor = max(0.01, 0.025 - (iteration - 1) * 0.005)
        fixed = _winsorize(fixed, limits=(kurt_winsor, kurt_winsor))
        applied.append(f"Winsorize {kurt_winsor*100:.1f}th pctile (kurtosis, iter {iteration})")

    if has_coll:
        fixed = _add_jitter(fixed, seed=seed + 10 + iteration, scale=jitter_scale)
        applied.append(f"Jitter (collinearity, iter {iteration})")

    # Rescale back to preserve interpretability
    fixed = _rescale_to_original(fixed, original)
    return fixed, applied

def _efa_fix_pass(df_current: pd.DataFrame, df_original: pd.DataFrame,
                  n_factors: int, rotation: str,
                  load_thresh: float, comm_thresh: float,
                  fix_history: dict, seed: int, iteration: int) -> tuple:
    """
    One complete fix pass:
      1. Run EFA on current df
      2. Identify all still-flagged variables (EFA diagnostics)
      3. Also detect raw data issues on those variables
      4. Apply targeted fixes
      5. Return updated df, updated diagnostics, still_flagged list
    """
    # Always ensure PSD before feeding to factor analyser
    df_safe = _ensure_psd_dataframe(df_current)
    fa_res  = run_efa(df_safe, n_factors, rotation)
    diag    = diagnose_loadings(fa_res["loadings"], fa_res["communalities"],
                                load_thresh, comm_thresh)
    still_flagged = diag[diag["RecommendDrop"]]["Variable"].tolist()

    if not still_flagged:
        return df_current, fa_res, diag, []

    df_next = df_current.copy()
    data_issues = _detect_data_issues(df_current)

    for var in still_flagged:
        original_series = df_original[var]
        current_series  = df_current[var]

        # Collect all known issues: raw data + EFA-level
        raw_tags = data_issues.get(var, [])

        # If raw data looks clean but EFA still flags it, synthesize issue tags
        # from the EFA diagnostic so we have something to act on
        if not raw_tags:
            skw = float(current_series.skew())
            krt = float(current_series.kurt())
            z   = (current_series - current_series.mean()) / max(current_series.std(), 1e-9)
            n_out = int((z.abs() > 2.5).sum())   # loosen outlier threshold for later iters

            if n_out > 0:
                raw_tags.append(f"outliers:{n_out}")
            if abs(skw) > 1.5:
                raw_tags.append(f"skewness:{skw:.3f}")
            if krt > 5:
                raw_tags.append(f"kurtosis:{krt:.3f}")
            # Fallback: low communality with no other handle → partial partialing
            if not raw_tags:
                # Orthogonalise vs. all other variables to boost unique variance
                others = [c for c in df_current.columns if c != var]
                residuals = current_series.copy()
                for ov in others:
                    cov = np.cov(residuals, df_current[ov])
                    if cov[1, 1] > 1e-9:
                        beta = cov[0, 1] / cov[1, 1]
                        residuals = residuals - beta * df_current[ov]
                # Blend: 70% residual + 30% original (keeps loadings but boosts uniqueness)
                blend = 0.30 * _rescale_to_original(residuals, current_series) + 0.70 * current_series
                df_next[var] = _rescale_to_original(blend, original_series)
                fix_history.setdefault(var, []).append(
                    f"Partial orthogonalisation blend (iter {iteration})")
                continue

        fixed_s, fix_desc_list = _apply_fixes_for_issues(
            current_series, raw_tags, original_series, seed=seed + iteration, iteration=iteration
        )
        df_next[var] = fixed_s
        fix_history.setdefault(var, []).extend(fix_desc_list)

    # Drop any zero-std columns that transforms may have created, replace with jitter
    for col in df_next.columns:
        if df_next[col].std() < 1e-9:
            df_next[col] = _add_jitter(df_original[col], seed=seed + hash(col) % 1000, scale=0.05)

    return df_next, fa_res, diag, still_flagged


def run_auto_fix(df: pd.DataFrame, initial_problem_vars: list,
                 n_factors: int, rotation: str,
                 load_thresh: float, comm_thresh: float,
                 seed: int = 42, max_iter: int = 6) -> tuple:
    """
    Iterative EFA-aware auto-fix engine.

    Loop:
      run EFA → find flagged vars → fix them → repeat
    until no variables are flagged or max_iter is reached.

    Returns
    -------
    df_fixed      : pd.DataFrame  — complete dataset, all variables retained
    fix_log       : pd.DataFrame  — per-variable log with issues + fixes applied
    final_efa     : dict          — EFA result on the final fixed dataset
    final_diag    : pd.DataFrame  — diagnostics on the final fixed dataset
    iteration_log : list[dict]    — summary of each iteration
    """
    df_current  = df.copy()
    df_original = df.copy()
    fix_history  = {}     # {var: [fix_desc, ...]}
    iteration_log = []

    for iteration in range(1, max_iter + 1):
        df_next, fa_res, diag, still_flagged = _efa_fix_pass(
            df_current, df_original,
            n_factors, rotation,
            load_thresh, comm_thresh,
            fix_history, seed, iteration,
        )

        n_flagged = len(still_flagged)
        avg_comm  = round(float(fa_res["communalities"].mean()), 4)
        iteration_log.append(dict(
            Iteration=iteration,
            FlaggedVars=n_flagged,
            AvgCommunality=avg_comm,
            FixedThisPass=", ".join(still_flagged) if still_flagged else "—",
        ))

        df_current = df_next

        if n_flagged == 0:
            break   # all clear — stop early

    # Final EFA on the fully fixed dataset (enforce PSD one last time)
    df_final_safe = _ensure_psd_dataframe(df_current)
    final_fa   = run_efa(df_final_safe, n_factors, rotation)
    final_diag = diagnose_loadings(final_fa["loadings"], final_fa["communalities"],
                                   load_thresh, comm_thresh)

    # Build per-variable fix log
    all_vars = df.columns.tolist()
    rows = []
    for var in all_vars:
        fixes = fix_history.get(var, [])
        data_issues_str = " | ".join(_detect_data_issues(df).get(var, ["None"]))
        if fixes:
            rows.append(dict(
                Variable=var,
                OriginalDataIssues=data_issues_str,
                FixesApplied=" → ".join(fixes),
                Iterations=len([f for f in fixes]),
                FinalCommunality=round(float(final_fa["communalities"][var]), 4),
                FinalIssue=final_diag.loc[final_diag["Variable"]==var,"Issue"].values[0]
                           if var in final_diag["Variable"].values else "—",
                Status="✔ Fixed" if var in fix_history else "✓ Clean",
            ))
        else:
            rows.append(dict(
                Variable=var,
                OriginalDataIssues=data_issues_str,
                FixesApplied="—",
                Iterations=0,
                FinalCommunality=round(float(final_fa["communalities"][var]), 4),
                FinalIssue=final_diag.loc[final_diag["Variable"]==var,"Issue"].values[0]
                           if var in final_diag["Variable"].values else "—",
                Status="✓ Clean",
            ))

    fix_log = pd.DataFrame(rows)
    return df_current, fix_log, final_fa, final_diag, pd.DataFrame(iteration_log)


# ══════════════════════════════════════════════════════════════════
# CFA FUNCTIONS  (core logic unchanged)
# ══════════════════════════════════════════════════════════════════

def build_cfa_model(loadings, threshold=0.4):
    factor_vars = {}
    for var in loadings.index:
        abs_row = np.abs(loadings.loc[var])
        if abs_row.max() >= threshold:
            best = abs_row.idxmax()
            factor_vars.setdefault(best, []).append(var)
    factor_vars = {f: v for f, v in factor_vars.items() if len(v) >= 2}
    model_str = "\n".join(f"{f} =~ " + " + ".join(v) for f, v in factor_vars.items())
    return model_str, factor_vars

def run_cfa(df, model_str):
    try:
        from semopy import Model
        from semopy.stats import calc_stats
        model = Model(model_str)
        model.fit(df)
        try:
            raw_stats = calc_stats(model)
            fit_indices = _parse_fit_indices(raw_stats)
        except Exception as e:
            fit_indices = {"parse_error": str(e)}
        return dict(success=True, fit_indices=fit_indices,
                    estimates=model.inspect(), model_str=model_str, error=None)
    except Exception as e:
        return dict(success=False, fit_indices={}, estimates=None, model_str=model_str, error=str(e))

def _parse_fit_indices(stats):
    flat = stats.iloc[0] if isinstance(stats, pd.DataFrame) else stats
    flat.index = [str(c).strip().upper() for c in flat.index]
    mapping = dict(CFI=["CFI"], TLI=["TLI","NNFI"], RMSEA=["RMSEA"], SRMR=["SRMR"],
                   Chi2=["CHI2","CHISQ","CHI-SQUARE","X2"], df=["DF","DOF"],
                   p_value=["P-VALUE","PVALUE","P_VALUE","P(CHI2)"], AIC=["AIC"], BIC=["BIC"])
    result = {}
    for label, candidates in mapping.items():
        for c in candidates:
            if c in flat.index:
                try: result[label] = round(float(flat[c]), 4)
                except: result[label] = flat[c]
                break
    return result

def assess_cfa_fit(fit_indices, thresholds):
    assessment = {}
    for idx, direction in [("CFI","≥"),("TLI","≥"),("RMSEA","≤"),("SRMR","≤")]:
        if idx in fit_indices and idx in thresholds:
            val, thresh = fit_indices[idx], thresholds[idx]
            passed = val >= thresh if direction == "≥" else val <= thresh
            assessment[idx] = dict(value=val, threshold=thresh, pass_=passed, direction=direction)
    n_pass = sum(1 for v in assessment.values() if v["pass_"])
    return dict(indices=assessment, n_pass=n_pass, n_total=len(assessment),
                overall_pass=(len(assessment)>0 and n_pass==len(assessment)))

def get_modification_suggestions(fit_assessment):
    suggestions, indices = [], fit_assessment.get("indices", {})
    checks = [
        ("RMSEA","RMSEA exceeds threshold. Consider freeing residual covariances between items sharing method variance, or removing items with high modification indices."),
        ("CFI",  "CFI is below threshold. Check whether indicators load on multiple factors. Consider adding cross-loadings or removing weak indicators."),
        ("SRMR", "SRMR exceeds threshold. Large residual correlations exist — review for systematic patterns among residuals."),
    ]
    for idx, msg in checks:
        if idx in indices and not indices[idx]["pass_"]:
            suggestions.append(f"{idx} = {indices[idx]['value']:.3f} — {msg}")
    if not suggestions and not fit_assessment.get("overall_pass"):
        suggestions.append("Model fit is inadequate. Consider revising factor structure, removing low-communality items, or allowing correlated residuals for items with common wording.")
    return suggestions


# ══════════════════════════════════════════════════════════════════
# SYNTHETIC DATA FUNCTIONS  (core logic unchanged)
# ══════════════════════════════════════════════════════════════════

def _make_psd(matrix, eps=1e-8):
    ev, evec = np.linalg.eigh(matrix)
    return evec @ np.diag(np.maximum(ev, eps)) @ evec.T

def generate_factor_based(df, efa_result, n_samples=500, seed=42):
    np.random.seed(seed)
    loadings = efa_result["loadings"].values
    communalities = efa_result["communalities"].values
    n_factors, columns = efa_result["n_factors"], efa_result["loadings"].index.tolist()
    factor_scores = np.random.multivariate_normal(np.zeros(n_factors), np.eye(n_factors), n_samples)
    common = factor_scores @ loadings.T
    unique = np.random.normal(0, np.sqrt(np.maximum(1-communalities, 1e-6)), (n_samples, len(columns)))
    synthetic = common + unique
    orig_mean, orig_std = df[columns].mean().values, df[columns].std().values
    orig_std = np.where(orig_std==0, 1.0, orig_std)
    syn_mean, syn_std = synthetic.mean(0), synthetic.std(0)
    syn_std = np.where(syn_std==0, 1.0, syn_std)
    rescaled = ((synthetic - syn_mean) / syn_std) * orig_std + orig_mean
    return pd.DataFrame(rescaled, columns=columns)

def generate_correlation_based(df, n_samples=500, seed=42):
    np.random.seed(seed)
    cov = _make_psd(df.cov().values)
    synthetic = np.random.multivariate_normal(df.mean().values, cov, n_samples)
    return pd.DataFrame(synthetic, columns=df.columns.tolist())

def validate_synthetic(original, synthetic):
    return pd.DataFrame([dict(
        Variable=col,
        OrigMean=round(original[col].mean(),4), SynMean=round(synthetic[col].mean(),4),
        OrigStd=round(original[col].std(),4),   SynStd=round(synthetic[col].std(),4),
        MeanDelta=round(abs(original[col].mean()-synthetic[col].mean()),4),
        StdDelta=round(abs(original[col].std()-synthetic[col].std()),4),
    ) for col in original.columns])


# ══════════════════════════════════════════════════════════════════
# PLOT FUNCTIONS  (core logic unchanged)
# ══════════════════════════════════════════════════════════════════

def plot_scree(eigenvalues, suggested_n):
    x = list(range(1, len(eigenvalues)+1))
    colors = [C["green"] if i < suggested_n else C["muted"] for i in range(len(eigenvalues))]
    fig = go.Figure()
    fig.add_trace(go.Bar(x=x, y=eigenvalues, marker_color=colors, name="Eigenvalue",
                         hovertemplate="Factor %{x}<br>λ = %{y:.3f}<extra></extra>"))
    fig.add_trace(go.Scatter(x=x, y=eigenvalues, mode="lines+markers",
                             line=dict(color=C["accent"], width=2),
                             marker=dict(size=7, color=C["accent"]), showlegend=False))
    fig.add_hline(y=1.0, line_dash="dash", line_color=C["yellow"],
                  annotation_text="Kaiser criterion (λ=1)", annotation_font_color=C["yellow"])
    fig.update_layout(**LAYOUT_BASE, height=360, showlegend=False,
                      title=dict(text=f"Scree Plot — Suggested factors: <b>{suggested_n}</b>",
                                 font=dict(color=C["accent2"])),
                      xaxis=dict(title="Factor", gridcolor=C["border"], tickmode="linear"),
                      yaxis=dict(title="Eigenvalue", gridcolor=C["border"]))
    return fig

def plot_loading_heatmap(loadings, threshold=0.4):
    z = loadings.values
    variables, factors = loadings.index.tolist(), loadings.columns.tolist()
    annotations = [dict(x=f, y=v, text=f"{z[i][j]:.2f}", showarrow=False,
                        font=dict(color="white" if abs(z[i][j])>=threshold else C["muted"], size=11))
                   for i, v in enumerate(variables) for j, f in enumerate(factors)]
    fig = go.Figure(go.Heatmap(z=z, x=factors, y=variables, zmid=0, zmin=-1, zmax=1,
                               colorscale=[[0,"#1e1b4b"],[0.5,C["surface"]],[1,C["accent"]]],
                               colorbar=dict(title="Loading", tickfont=dict(color=C["text"]),
                                             title_font=dict(color=C["text"])),
                               hovertemplate="%{y} → %{x}<br>Loading: %{z:.3f}<extra></extra>"))
    fig.update_layout(**LAYOUT_BASE, annotations=annotations,
                      title=dict(text="Factor Loading Heatmap", font=dict(color=C["accent2"])),
                      height=max(300, len(variables)*32+100), xaxis=dict(side="top"))
    return fig

def plot_communalities(communalities, comm_thresh=0.3):
    colors = [C["green"] if v>=0.5 else (C["yellow"] if v>=comm_thresh else C["red"])
              for v in communalities.values]
    fig = go.Figure(go.Bar(x=communalities.index.tolist(), y=communalities.values,
                           marker_color=colors,
                           hovertemplate="%{x}<br>Communality: %{y:.3f}<extra></extra>"))
    fig.add_hline(y=comm_thresh, line_dash="dash", line_color=C["red"],
                  annotation_text=f"Threshold ({comm_thresh})", annotation_font_color=C["red"])
    fig.add_hline(y=0.5, line_dash="dot", line_color=C["yellow"],
                  annotation_text="Good (0.50)", annotation_font_color=C["yellow"])
    fig.update_layout(**LAYOUT_BASE, height=360,
                      title=dict(text="Communalities per Variable", font=dict(color=C["accent2"])),
                      xaxis=dict(title="Variable", gridcolor=C["border"], tickangle=-35),
                      yaxis=dict(title="Communality", gridcolor=C["border"], range=[0,1]))
    return fig

def plot_fit_indices(fit_assessment):
    indices = fit_assessment.get("indices", {})
    if not indices:
        return go.Figure()
    labels = list(indices.keys())
    values = [d["value"] for d in indices.values()]
    colors = [C["green"] if d["pass_"] else C["red"] for d in indices.values()]
    fig = go.Figure(go.Bar(x=labels, y=values, marker_color=colors,
                           text=[f"{v:.3f}" for v in values], textposition="outside",
                           hovertemplate="%{x}: %{y:.4f}<extra></extra>"))
    for label, data in indices.items():
        fig.add_annotation(x=label, y=data["threshold"],
                           text=f"Threshold: {data['threshold']}",
                           showarrow=True, arrowhead=2, arrowcolor=C["yellow"],
                           font=dict(color=C["yellow"], size=10), ay=-30)
    fig.update_layout(**LAYOUT_BASE, height=360, showlegend=False,
                      title=dict(text="CFA Fit Indices", font=dict(color=C["accent2"])),
                      xaxis=dict(title="Index", gridcolor=C["border"]),
                      yaxis=dict(title="Value", gridcolor=C["border"]))
    return fig

def plot_correlation_matrix(df):
    corr = df.corr(numeric_only=True).round(2)
    cols = corr.columns.tolist()
    annotations = [dict(x=c, y=r, text=f"{corr.loc[r,c]:.2f}", showarrow=False,
                        font=dict(size=9, color="white" if abs(corr.loc[r,c])>0.5 else C["text"]))
                   for r in corr.index for c in cols]
    fig = go.Figure(go.Heatmap(z=corr.values, x=cols, y=cols, colorscale="RdBu", zmid=0, zmin=-1, zmax=1,
                               colorbar=dict(title="r", tickfont=dict(color=C["text"]),
                                             title_font=dict(color=C["text"])),
                               hovertemplate="%{y} × %{x}<br>r = %{z:.2f}<extra></extra>"))
    fig.update_layout(**LAYOUT_BASE, annotations=annotations,
                      title=dict(text="Correlation Matrix", font=dict(color=C["accent2"])),
                      height=max(350, len(cols)*30+120), xaxis=dict(tickangle=-35))
    return fig

def plot_synthetic_comparison(original, synthetic, max_vars=6):
    cols = original.columns[:max_vars].tolist()
    ncols = min(3, len(cols))
    nrows = (len(cols)+ncols-1)//ncols
    fig = make_subplots(rows=nrows, cols=ncols, subplot_titles=cols)
    for i, col in enumerate(cols):
        r, ci = i//ncols+1, i%ncols+1
        fig.add_trace(go.Histogram(x=original[col], name="Original", nbinsx=20,
                                   marker_color=C["accent"], opacity=0.6, showlegend=(i==0)), row=r, col=ci)
        fig.add_trace(go.Histogram(x=synthetic[col], name="Synthetic", nbinsx=20,
                                   marker_color=C["accent2"], opacity=0.6, showlegend=(i==0)), row=r, col=ci)
    fig.update_layout(**LAYOUT_BASE, barmode="overlay", height=300*nrows,
                      title=dict(text="Original vs Synthetic Distributions", font=dict(color=C["accent2"])),
                      legend=dict(orientation="h", y=1.02, x=0.5, xanchor="center", font=dict(color=C["text"])))
    return fig


# ══════════════════════════════════════════════════════════════════
# DOCX REPORT GENERATOR
# ══════════════════════════════════════════════════════════════════

def generate_docx_report(original_df, cleaned_df, suitability, efa_result,
                         diagnostics, dropped_vars, cfa_result, fit_assessment,
                         cfa_thresholds, synthetic_validation=None, model_str=""):
    """Generate a Word (.docx) analysis report and return bytes."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = DocxDocument()

    # ── Page margins ────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Helper styles ───────────────────────────────────────────────
    def _set_run_color(run, hex_color):
        run.font.color.rgb = RGBColor(
            int(hex_color[1:3], 16),
            int(hex_color[3:5], 16),
            int(hex_color[5:7], 16),
        )

    def _add_heading(text, level=1, color="#6c8dfa"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16 if level == 1 else 13)
        _set_run_color(run, color)
        return p

    def _add_para(text, bold=False, color=None, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            _set_run_color(run, color)
        return p

    def _shade_cell(cell, hex_color):
        """Apply background shading to a table cell."""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.lstrip("#"))
        tc_pr.append(shd)

    def _add_table(headers, rows_data, col_widths=None, header_bg="1e2540"):
        """Add a table with styled header row."""
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        # Header row
        hdr_cells = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            _set_run_color(run, "#6c8dfa")
            _shade_cell(hdr_cells[i], header_bg)
        # Data rows
        for row_vals, row_colors in rows_data:
            cells = t.add_row().cells
            for i, val in enumerate(row_vals):
                cells[i].text = str(val)
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(9)
                if row_colors and row_colors[i]:
                    _set_run_color(run, row_colors[i])
        return t

    ts = datetime.now().strftime("%Y-%m-%d %H:%M")

    # ── Title ────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("◈ EFActor — Analysis Report")
    r.bold = True
    r.font.size = Pt(20)
    _set_run_color(r, "#6c8dfa")

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_p.add_run(f"Generated: {ts}")
    mr.font.size = Pt(9)
    _set_run_color(mr, "#94a3b8")

    doc.add_paragraph()

    # ── 1. Dataset Overview ──────────────────────────────────────────
    _add_heading("1. Dataset Overview", level=1)
    overview_rows = [
        (["Original Variables", str(len(original_df.columns))], [None, "#6c8dfa"]),
        (["After Cleaning",     str(len(cleaned_df.columns))],  [None, "#6c8dfa"]),
        (["Observations",       str(len(cleaned_df))],           [None, "#6c8dfa"]),
        (["Dropped Variables",  str(len(dropped_vars))],         [None, "#f87171" if dropped_vars else "#34d399"]),
    ]
    _add_table(["Metric", "Value"], overview_rows)
    if dropped_vars:
        _add_para(f"Dropped: {', '.join(dropped_vars)}", color="#f87171", size=9)

    # ── 2. EFA Suitability ───────────────────────────────────────────
    _add_heading("2. EFA Suitability", level=1)
    s = suitability
    suit_rows = [
        (["KMO Score",           f"{s['kmo_model']} — {s['kmo_label']}", "≥ 0.60",
          "PASS" if s["kmo_pass"] else "FAIL"],
         [None, None, None, "#34d399" if s["kmo_pass"] else "#f87171"]),
        (["Bartlett's Sphericity", f"χ² = {s['bartlett_chi2']}, p = {s['bartlett_p']}", "p < 0.05",
          "PASS" if s["bartlett_pass"] else "FAIL"],
         [None, None, None, "#34d399" if s["bartlett_pass"] else "#f87171"]),
        (["Overall",             "",  "",
          "PASS" if s["overall_pass"] else "FAIL"],
         [None, None, None, "#34d399" if s["overall_pass"] else "#f87171"]),
    ]
    _add_table(["Test", "Value", "Threshold", "Result"], suit_rows)

    # ── 3. EFA Results ───────────────────────────────────────────────
    _add_heading("3. Exploratory Factor Analysis", level=1)
    _add_para(f"Factors extracted: {efa_result['n_factors']}", bold=True, color="#6c8dfa")

    loadings      = efa_result["loadings"]
    communalities = efa_result["communalities"]
    variance      = efa_result["variance"]

    # Loadings table
    _add_heading("Factor Loadings", level=2, color="#a78bfa")
    load_headers = ["Variable"] + loadings.columns.tolist() + ["Communality"]
    load_rows = []
    for v in loadings.index:
        comm_val = communalities[v]
        comm_color = "#34d399" if comm_val >= 0.5 else ("#fbbf24" if comm_val >= 0.3 else "#f87171")
        vals   = [v] + [f"{loadings.loc[v, c]:.3f}" for c in loadings.columns] + [f"{comm_val:.3f}"]
        colors = [None] + [None] * len(loadings.columns) + [comm_color]
        load_rows.append((vals, colors))
    _add_table(load_headers, load_rows)

    # Variance table
    _add_heading("Variance Explained", level=2, color="#a78bfa")
    var_rows = [
        ([idx, f"{row['SS Loadings']:.4f}",
          f"{row['Proportion Var']*100:.2f}%",
          f"{row['Cumulative Var']*100:.2f}%"], [None]*4)
        for idx, row in variance.iterrows()
    ]
    _add_table(["Factor", "SS Loadings", "Proportion Var", "Cumulative Var"], var_rows)

    # Diagnostics table
    _add_heading("Item Diagnostics", level=2, color="#a78bfa")
    diag_rows = []
    for _, r in diagnostics.iterrows():
        comm = r["Communality"]
        comm_color = "#34d399" if comm >= 0.5 else ("#fbbf24" if comm >= 0.3 else "#f87171")
        issue      = r["Issue"]
        issue_color = "#34d399" if issue == "OK" else ("#fbbf24" if "Cross" in issue else "#f87171")
        drop_color  = "#f87171" if r["RecommendDrop"] else "#34d399"
        drop_text   = "✗ Drop" if r["RecommendDrop"] else "✓ Keep"
        diag_rows.append((
            [r["Variable"], str(r["MaxLoading"]), str(r["FactorsAboveThreshold"]),
             str(comm), issue, drop_text],
            [None, None, None, comm_color, issue_color, drop_color]
        ))
    _add_table(["Variable", "Max Load", "# Factors ≥ Threshold", "Communality", "Issue", "Recommended"],
               diag_rows)

    # ── 4. CFA ──────────────────────────────────────────────────────
    if cfa_result and cfa_result.get("success") and fit_assessment:
        _add_heading("4. Confirmatory Factor Analysis", level=1)
        fa = fit_assessment
        overall_text  = "ADEQUATE" if fa["overall_pass"] else "INADEQUATE"
        overall_color = "#34d399"  if fa["overall_pass"] else "#f87171"
        _add_para(f"Model Fit: {overall_text}  ({fa['n_pass']}/{fa['n_total']} indices passed)",
                  bold=True, color=overall_color)
        fit_rows = [
            ([idx, str(d["value"]), f"{d['direction']} {d['threshold']}",
              "PASS" if d["pass_"] else "FAIL"],
             [None, None, None, "#34d399" if d["pass_"] else "#f87171"])
            for idx, d in fa["indices"].items()
        ]
        _add_table(["Index", "Value", "Threshold", "Result"], fit_rows)

        _add_heading("Model Specification", level=2, color="#a78bfa")
        doc.add_paragraph(model_str or "—")

    # ── 5. Synthetic Data Validation ────────────────────────────────
    if synthetic_validation is not None:
        _add_heading("5. Synthetic Data Validation", level=1)
        syn_rows = [
            ([r["Variable"], str(r["OrigMean"]), str(r["SynMean"]),
              str(r["OrigStd"]),  str(r["SynStd"]),
              str(r["MeanDelta"]), str(r["StdDelta"])], [None]*7)
            for _, r in synthetic_validation.iterrows()
        ]
        _add_table(["Variable", "Orig Mean", "Syn Mean", "Orig Std",
                    "Syn Std", "Mean Δ", "Std Δ"], syn_rows)

    # ── Footer ───────────────────────────────────────────────────────
    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot_p.add_run(
        "EFActor — Results should be interpreted in context of your research design "
        "and theoretical framework."
    )
    fr.font.size = Pt(8)
    _set_run_color(fr, "#64748b")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════
# SESSION STATE INIT
# ══════════════════════════════════════════════════════════════════
_DEFAULTS = dict(df_original=None, df_working=None, suitability=None,
                 n_factors_auto=None, eigenvalues=None, efa_result=None,
                 diagnostics=None, efa_done=False, cfa_result=None,
                 fit_assessment=None, synthetic_factor=None, synthetic_corr=None,
                 syn_validation=None, report_docx=None,
                 df_autofix=None, fix_log=None, autofix_efa_result=None)
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v
if "dropped_vars" not in st.session_state:
    st.session_state["dropped_vars"] = []
S = st.session_state

is_trial   = S.get("is_free_trial", False)
credits    = S.get("credits", 0)
key_owner  = S.get("key_owner", "Researcher")
access_key = S.get("access_key", "")


# ══════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════
with st.sidebar:
    cred_color   = "#34d399" if (is_trial or credits > 10) else ("#fbbf24" if credits > 3 else "#f87171")
    cred_display = "Trial" if is_trial else str(credits)
    cred_label   = "Free Trial — export locked" if is_trial else "Credits remaining"
    cred_bar_w   = "100" if is_trial else str(min(100, int(credits * 2)))

    st.markdown(f"""
    <div style="padding:16px 0 20px;border-bottom:1px solid #1e2540;margin-bottom:20px;">
      <div style="display:flex;align-items:center;gap:10px;margin-bottom:14px;">
        <div style="width:32px;height:32px;background:linear-gradient(135deg,#6c8dfa,#a78bfa);
                    border-radius:8px;display:flex;align-items:center;justify-content:center;
                    font-size:16px;flex-shrink:0;">🔬</div>
        <div>
          <div style="font-family:Inter,sans-serif;font-size:16px;font-weight:800;
                      background:linear-gradient(90deg,#6c8dfa,#a78bfa);
                      -webkit-background-clip:text;-webkit-text-fill-color:transparent;">EFActor</div>
          <div style="font-family:Inter,sans-serif;font-size:9px;color:#64748b;
                      letter-spacing:2px;text-transform:uppercase;font-weight:500;">Psychometric Analysis</div>
        </div>
      </div>
      <div style="background:#0a0c14;border:1px solid #1e2540;border-radius:10px;padding:12px 14px;">
        <div style="font-family:Inter,sans-serif;font-size:9px;color:#64748b;
                    letter-spacing:2px;text-transform:uppercase;font-weight:500;margin-bottom:6px;">Account</div>
        <div style="font-family:Inter,sans-serif;font-size:13px;font-weight:600;
                    color:#e2e8f0;margin-bottom:8px;">{key_owner}</div>
        <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:6px;">
          <div style="font-family:Inter,sans-serif;font-size:10px;color:#64748b;">{cred_label}</div>
          <div style="font-family:Inter,sans-serif;font-size:18px;font-weight:800;
                      color:{cred_color};">{cred_display}</div>
        </div>
        <div style="background:#1e2540;border-radius:4px;height:3px;overflow:hidden;">
          <div style="background:{cred_color};height:3px;width:{cred_bar_w}%;border-radius:4px;"></div>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Sign Out button
    if st.button("↩ Sign Out", use_container_width=True, key="signout_btn"):
        for k in ["authenticated","access_key","key_owner","credits","is_free_trial",
                  "_last_file_key","df_original","df_working","suitability","n_factors_auto",
                  "eigenvalues","efa_result","diagnostics","efa_done","dropped_vars",
                  "cfa_result","fit_assessment","synthetic_factor","synthetic_corr",
                  "syn_validation","report_docx","df_autofix","fix_log","autofix_efa_result"]:
            st.session_state.pop(k, None)
        st.rerun()

    st.markdown("<div style='margin-bottom:6px;'></div>", unsafe_allow_html=True)

    st.markdown("### ⚙️ EFA Settings")
    loading_threshold    = st.slider("Loading threshold",    0.30, 0.60, 0.40, 0.05)
    communality_threshold= st.slider("Communality threshold",0.20, 0.60, 0.30, 0.05)
    rotation_method      = st.selectbox("Rotation method",
                             ["varimax","oblimin","promax","quartimax","equamax"])

    st.divider()
    st.markdown("### 📐 CFA Thresholds")
    cfi_thresh   = st.slider("CFI ≥",   0.80, 0.99, 0.95, 0.01)
    tli_thresh   = st.slider("TLI ≥",   0.80, 0.99, 0.95, 0.01)
    rmsea_thresh = st.slider("RMSEA ≤", 0.04, 0.15, 0.06, 0.01)
    srmr_thresh  = st.slider("SRMR ≤",  0.04, 0.15, 0.08, 0.01)
    cfa_thresholds = dict(CFI=cfi_thresh, TLI=tli_thresh, RMSEA=rmsea_thresh, SRMR=srmr_thresh)

    st.divider()
    st.markdown("### 🧪 Synthetic Data")
    syn_n    = st.number_input("Sample size",  50, 10000, 500, 50)
    syn_seed = st.number_input("Random seed",   0,  9999,  42)

    st.divider()
    if st.button("🔄 Reset Analysis", use_container_width=True):
        for k in ["_last_file_key","df_original","df_working","suitability","n_factors_auto",
                  "eigenvalues","efa_result","diagnostics","efa_done","dropped_vars",
                  "cfa_result","fit_assessment","synthetic_factor","synthetic_corr",
                  "syn_validation","report_docx","df_autofix","fix_log","autofix_efa_result"]:
            st.session_state.pop(k, None)
        st.rerun()

    if is_trial:
        st.divider()
        st.markdown("""
        <div style="background:rgba(108,141,250,.07);border:1px solid rgba(108,141,250,.2);
                    border-radius:10px;padding:14px 16px;text-align:center;">
          <div style="font-family:Inter,sans-serif;font-size:12px;font-weight:700;
                      color:#6c8dfa;margin-bottom:6px;">Unlock Exports</div>
          <div style="font-family:Inter,sans-serif;font-size:11px;color:#64748b;line-height:1.6;margin-bottom:10px;">
            Get credits to download datasets and reports.
          </div>
        </div>
        """, unsafe_allow_html=True)
        st.link_button("Get Access Key →", "https://your-payment-link", use_container_width=True)


# ══════════════════════════════════════════════════════════════════
# MAIN HEADER
# ══════════════════════════════════════════════════════════════════
st.markdown("""
<div style="display:flex;align-items:center;gap:14px;margin-bottom:6px;">
  <div style="width:40px;height:40px;background:linear-gradient(135deg,#6c8dfa,#a78bfa);
              border-radius:10px;display:flex;align-items:center;justify-content:center;
              font-size:20px;flex-shrink:0;">🔬</div>
  <div>
    <div style="font-family:Inter,sans-serif;font-size:28px;font-weight:800;
                background:linear-gradient(90deg,#6c8dfa,#a78bfa);
                -webkit-background-clip:text;-webkit-text-fill-color:transparent;
                letter-spacing:-0.5px;line-height:1;">EFActor</div>
    <div style="font-family:Inter,sans-serif;font-size:11px;color:#64748b;
                letter-spacing:2px;text-transform:uppercase;font-weight:500;margin-top:2px;">
      Psychometric Analysis Platform
    </div>
  </div>
</div>
""", unsafe_allow_html=True)
st.markdown("_Upload your dataset and run a complete EFA → CFA pipeline. Diagnose items, confirm structure, generate synthetic data, and export your results._")

if is_trial:
    st.markdown('<div class="warn-box">⚡ <b>Free Trial mode</b> — all analysis runs are unlimited. Data export requires an access key with credits. <a href="https://x.com/bayantx360" style="color:#fbbf24;">Get access →</a></div>', unsafe_allow_html=True)

st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# STEP 1 — UPLOAD
# ══════════════════════════════════════════════════════════════════
st.markdown('<div class="step-badge"><span class="step-num">1</span> Upload Dataset</div>', unsafe_allow_html=True)

col_up1, col_up2 = st.columns([2, 1])
with col_up1:
    uploaded = st.file_uploader("Upload CSV or Excel",
                                 type=["csv","CSV","xlsx","xls","XLSX","XLS"],
                                 help="Numeric columns only. Rows with missing values are auto-dropped.")
with col_up2:
    st.markdown('<div class="info-box">💡 <b>Requirements</b><br>• Numeric variables only<br>• Minimum 5 variables<br>• Recommended ≥ 100 rows<br>• Missing values auto-dropped</div>', unsafe_allow_html=True)

if uploaded:
    _file_key = f"{uploaded.name}_{uploaded.size}"
    if S.get("_last_file_key") != _file_key:
        try:
            fname = uploaded.name.lower()
            if fname.endswith(".csv"):
                try:
                    df_raw = pd.read_csv(uploaded, encoding="utf-8")
                except UnicodeDecodeError:
                    uploaded.seek(0)
                    df_raw = pd.read_csv(uploaded, encoding="latin-1")
            elif fname.endswith((".xlsx", ".xls")):
                df_raw = pd.read_excel(uploaded)
            else:
                st.error("❌ Unsupported file type.")
                st.stop()
            # Fix: coerce string-encoded values, drop all-NaN cols, then drop NaN rows
            df_coerced = df_raw.apply(lambda col: pd.to_numeric(col, errors="coerce"))
            df_numeric = (df_coerced
                          .select_dtypes(include=[np.number])
                          .dropna(axis=1, how="all")
                          .dropna())
            if len(df_numeric.columns) < 3:
                st.error("❌ Need at least 3 numeric variables. Non-numeric columns are excluded automatically.")
                st.stop()
            if len(df_numeric) < 30:
                st.warning("⚠️ Fewer than 30 observations — results may be unreliable.")
            S.df_original = df_numeric.copy()
            S.df_working  = df_numeric.copy()
            S.dropped_vars = []
            for k in ["suitability","efa_result","cfa_result","fit_assessment","efa_done",
                      "synthetic_factor","synthetic_corr","syn_validation","report_docx",
                      "n_factors_auto","eigenvalues","diagnostics"]:
                S[k] = None if k != "efa_done" else False
            S["_last_file_key"] = _file_key
        except Exception as e:
            st.error(f"❌ Could not parse file: {e}")
            st.stop()

if S.df_original is None:
    st.markdown('<div class="warn-box">👆 Upload a dataset to begin.</div>', unsafe_allow_html=True)
    st.stop()

with st.expander("📋 Dataset Preview", expanded=True):
    c1, c2, c3, c4 = st.columns(4)
    for col, val, label in [(c1,len(S.df_original),"Observations"),(c2,len(S.df_original.columns),"Variables"),
                             (c3,len(S.df_working.columns),"Working Vars"),(c4,len(S.dropped_vars),"Dropped")]:
        col.markdown(f'<div class="metric-card"><div class="metric-val">{val}</div><div class="metric-label">{label}</div></div>', unsafe_allow_html=True)
    st.dataframe(S.df_working.head(10), use_container_width=True)
    t1, t2 = st.tabs(["Descriptive Statistics", "Correlation Matrix"])
    with t1:
        st.dataframe(S.df_working.describe().T.round(4), use_container_width=True)
    with t2:
        st.plotly_chart(plot_correlation_matrix(S.df_working), use_container_width=True)

st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# STEP 2 — EFA SUITABILITY
# ══════════════════════════════════════════════════════════════════
st.markdown('<div class="step-badge"><span class="step-num">2</span> EFA Suitability Tests</div>', unsafe_allow_html=True)

if st.button("▶ Run Suitability Tests"):
    with st.spinner("Running KMO and Bartlett's tests…"):
        S.suitability = check_efa_suitability(S.df_working)
        ev_info = determine_n_factors(S.df_working)
        S.n_factors_auto = ev_info["suggested_n"]
        S.eigenvalues = ev_info["eigenvalues"]

if S.suitability:
    suit = S.suitability
    pill = "pill-pass" if suit["overall_pass"] else "pill-fail"
    text = "✓ SUITABLE FOR EFA" if suit["overall_pass"] else "✗ EFA SUITABILITY ISSUES"
    st.markdown(f'<b>Overall:</b> <span class="{pill}">{text}</span>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    c1.markdown(f'<div class="metric-card"><div class="metric-val">{suit["kmo_model"]}</div>'
                f'<div class="metric-label">KMO Score — {suit["kmo_label"]}</div>'
                f'<div style="margin-top:8px"><span class="{"pill-pass" if suit["kmo_pass"] else "pill-fail"}">{"PASS" if suit["kmo_pass"] else "FAIL"}</span></div></div>',
                unsafe_allow_html=True)
    c2.markdown(f'<div class="metric-card"><div class="metric-val">{suit["bartlett_p"]}</div>'
                f'<div class="metric-label">Bartlett p-value (χ²={suit["bartlett_chi2"]})</div>'
                f'<div style="margin-top:8px"><span class="{"pill-pass" if suit["bartlett_pass"] else "pill-fail"}">{"PASS" if suit["bartlett_pass"] else "FAIL"}</span></div></div>',
                unsafe_allow_html=True)
    if not suit["overall_pass"]:
        st.markdown('<div class="warn-box">⚠️ Suitability concerns detected. You can still proceed — interpret results cautiously.</div>', unsafe_allow_html=True)
    st.markdown("#### Scree Plot")
    st.plotly_chart(plot_scree(S.eigenvalues, S.n_factors_auto), use_container_width=True)
    st.markdown(f'<div class="info-box">🔢 Kaiser criterion suggests <b>{S.n_factors_auto}</b> factor(s). Override below if theory or scree plot suggests otherwise.</div>', unsafe_allow_html=True)

st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# STEP 3 — EFA
# ══════════════════════════════════════════════════════════════════
st.markdown('<div class="step-badge"><span class="step-num">3</span> Exploratory Factor Analysis (EFA)</div>', unsafe_allow_html=True)

if S.suitability is None:
    st.markdown('<div class="warn-box">Complete Step 2 first.</div>', unsafe_allow_html=True)
else:
    _max_factors = max(2, len(S.df_working.columns) - 1)
    default_n = S.n_factors_auto if S.n_factors_auto else 2
    default_n = int(min(max(1, default_n), _max_factors))  # clamp to valid range
    n_factors = st.number_input("Number of factors to extract", 1, _max_factors, default_n, 1)
    if st.button("▶ Run EFA", use_container_width=True):
        with st.spinner("Running factor analysis…"):
            S.efa_result  = run_efa(S.df_working, n_factors, rotation_method)
            S.diagnostics = diagnose_loadings(S.efa_result["loadings"],
                                              S.efa_result["communalities"],
                                              loading_threshold, communality_threshold)
            S.efa_done = True
    if S.efa_result:
        t1, t2, t3, t4 = st.tabs(["Factor Loadings","Variance Explained","Loading Heatmap","Communalities"])
        with t1: st.dataframe(S.efa_result["loadings"].round(4), use_container_width=True)
        with t2: st.dataframe(S.efa_result["variance"].round(4), use_container_width=True)
        with t3: st.plotly_chart(plot_loading_heatmap(S.efa_result["loadings"], loading_threshold), use_container_width=True)
        with t4: st.plotly_chart(plot_communalities(S.efa_result["communalities"], communality_threshold), use_container_width=True)

st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# STEP 4 — DIAGNOSE, AUTO-FIX & REFINE
# ══════════════════════════════════════════════════════════════════
st.markdown('<div class="step-badge"><span class="step-num">4</span> Diagnose, Auto-Fix &amp; Refine Items</div>', unsafe_allow_html=True)

if S.efa_result is None:
    st.markdown('<div class="warn-box">Complete EFA (Step 3) first.</div>', unsafe_allow_html=True)
else:
    # ── Colour-coded diagnostics table ──────────────────────────────
    def _diag_table_html(df):
        rows_html = ""
        for _, r in df.iterrows():
            # Communality colouring
            comm = r["Communality"]
            if comm >= 0.5:
                comm_style = "color:#34d399;font-weight:600;"
                comm_icon  = "🟢"
            elif comm >= 0.3:
                comm_style = "color:#fbbf24;font-weight:600;"
                comm_icon  = "🟡"
            else:
                comm_style = "color:#f87171;font-weight:600;"
                comm_icon  = "🔴"

            # Issue colouring
            issue = r["Issue"]
            if issue == "OK":
                issue_style = "color:#34d399;font-weight:600;"
                issue_icon  = "✓ OK"
            elif "Cross-Loader" in issue:
                issue_style = "color:#fbbf24;font-weight:600;"
                issue_icon  = f"⚠ {issue}"
            else:
                issue_style = "color:#f87171;font-weight:600;"
                issue_icon  = f"✗ {issue}"

            # RecommendDrop colouring
            drop = r["RecommendDrop"]
            if drop:
                drop_style = "color:#f87171;font-weight:700;"
                drop_cell  = "✗ Drop"
            else:
                drop_style = "color:#34d399;font-weight:600;"
                drop_cell  = "✓ Keep"

            rows_html += (
                f"<tr>"
                f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;'>{r['Variable']}</td>"
                f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;'>{r['MaxLoading']}</td>"
                f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;text-align:center;'>{r['FactorsAboveThreshold']}</td>"
                f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;{comm_style}'>{comm_icon} {comm}</td>"
                f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;{issue_style}'>{issue_icon}</td>"
                f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;{drop_style}'>{drop_cell}</td>"
                f"</tr>"
            )
        return f"""
<div style="overflow-x:auto;">
<table style="width:100%;border-collapse:collapse;font-size:.85rem;font-family:Inter,sans-serif;">
<thead>
<tr style="background:#1e2540;">
  <th style="padding:8px 10px;text-align:left;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Variable</th>
  <th style="padding:8px 10px;text-align:left;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Max Loading</th>
  <th style="padding:8px 10px;text-align:center;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;"># Factors ≥ Threshold</th>
  <th style="padding:8px 10px;text-align:left;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Communality</th>
  <th style="padding:8px 10px;text-align:left;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Issue</th>
  <th style="padding:8px 10px;text-align:left;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Recommended</th>
</tr>
</thead>
<tbody style="color:#e2e8f0;">
{rows_html}
</tbody>
</table>
</div>
"""
    st.markdown(_diag_table_html(S.diagnostics), unsafe_allow_html=True)
    problem_vars = S.diagnostics[S.diagnostics["RecommendDrop"]]["Variable"].tolist()
    if problem_vars:
        st.markdown(f'<div class="warn-box">⚠️ <b>{len(problem_vars)}</b> variable(s) flagged: {", ".join(problem_vars)}</div>', unsafe_allow_html=True)

    # ── AUTO-FIX SECTION ────────────────────────────────────────────
    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)
    st.markdown("""
    <div style="background:rgba(52,211,153,.06);border:1px solid rgba(52,211,153,.25);
                border-left:3px solid #34d399;border-radius:8px;padding:16px 20px;margin-bottom:18px;">
      <div style="font-family:Inter,sans-serif;font-size:15px;font-weight:700;color:#34d399;margin-bottom:6px;">
        ⚡ EFA Auto-Fix Engine
      </div>
      <div style="font-family:Inter,sans-serif;font-size:13px;color:#94a3b8;line-height:1.7;">
        Automatically detects data-level causes of EFA problems (outliers, skewness, kurtosis,
        near-zero variance, collinearity) and applies targeted fixes — <b style="color:#e2e8f0;">without dropping any variable</b>.
        The complete fixed dataset is returned and can be used for downstream EFA / CFA.
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style="display:grid;grid-template-columns:repeat(5,1fr);gap:10px;margin-bottom:18px;">
      <div style="background:#111520;border:1px solid #1e2540;border-radius:8px;padding:12px 10px;text-align:center;">
        <div style="font-size:18px;margin-bottom:6px;">🎯</div>
        <div style="font-family:Inter,sans-serif;font-size:11px;font-weight:600;color:#6c8dfa;margin-bottom:4px;">Winsorize</div>
        <div style="font-family:Inter,sans-serif;font-size:10px;color:#64748b;line-height:1.5;">Clips extreme outliers at 5th–95th pctile</div>
      </div>
      <div style="background:#111520;border:1px solid #1e2540;border-radius:8px;padding:12px 10px;text-align:center;">
        <div style="font-size:18px;margin-bottom:6px;">📐</div>
        <div style="font-family:Inter,sans-serif;font-size:11px;font-weight:600;color:#a78bfa;margin-bottom:4px;">Log / √ Transform</div>
        <div style="font-family:Inter,sans-serif;font-size:10px;color:#64748b;line-height:1.5;">Corrects high skewness</div>
      </div>
      <div style="background:#111520;border:1px solid #1e2540;border-radius:8px;padding:12px 10px;text-align:center;">
        <div style="font-size:18px;margin-bottom:6px;">🔀</div>
        <div style="font-family:Inter,sans-serif;font-size:11px;font-weight:600;color:#34d399;margin-bottom:4px;">Jitter</div>
        <div style="font-family:Inter,sans-serif;font-size:10px;color:#64748b;line-height:1.5;">Resolves zero-variance & collinearity</div>
      </div>
      <div style="background:#111520;border:1px solid #1e2540;border-radius:8px;padding:12px 10px;text-align:center;">
        <div style="font-size:18px;margin-bottom:6px;">📊</div>
        <div style="font-family:Inter,sans-serif;font-size:11px;font-weight:600;color:#fbbf24;margin-bottom:4px;">Kurtosis Fix</div>
        <div style="font-family:Inter,sans-serif;font-size:10px;color:#64748b;line-height:1.5;">Winsorize at 2.5th–97.5th for peaked dists</div>
      </div>
      <div style="background:#111520;border:1px solid #1e2540;border-radius:8px;padding:12px 10px;text-align:center;">
        <div style="font-size:18px;margin-bottom:6px;">↔</div>
        <div style="font-family:Inter,sans-serif;font-size:11px;font-weight:600;color:#f87171;margin-bottom:4px;">Rescale</div>
        <div style="font-family:Inter,sans-serif;font-size:10px;color:#64748b;line-height:1.5;">Restores original mean & SD after transform</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    col_fix1, col_fix2 = st.columns([2, 1])
    with col_fix1:
        if st.button("⚡ Run Auto-Fix (keep all variables)", use_container_width=True, key="autofix_btn"):
            with st.spinner("Running iterative EFA-aware fix engine (up to 6 passes)…"):
                df_fixed, fix_log_df, final_fa, final_diag, iter_log_df = run_auto_fix(
                    S.df_working,
                    initial_problem_vars=problem_vars,
                    n_factors=S.efa_result["n_factors"],
                    rotation=rotation_method,
                    load_thresh=loading_threshold,
                    comm_thresh=communality_threshold,
                    seed=42,
                    max_iter=6,
                )
                S.df_autofix         = df_fixed
                S.fix_log            = fix_log_df
                S.autofix_efa_result = dict(efa=final_fa, diag=final_diag, iter_log=iter_log_df)
            n_remaining = int(final_diag["RecommendDrop"].sum())
            if n_remaining == 0:
                st.success(f"✓ All variables fixed. All {len(df_fixed.columns)} variables pass EFA diagnostics.")
            else:
                st.warning(f"⚠️ {n_remaining} variable(s) still flagged after 6 iterations — these have structural issues that cannot be resolved by data transformation alone. The best-possible fixed dataset is still returned with all variables retained.")

    with col_fix2:
        if S.df_autofix is not None:
            if st.button("✅ Use Fixed Dataset for CFA / Export", use_container_width=True, key="apply_fix_btn"):
                S.df_working = S.df_autofix.copy()
                # Re-run EFA on working set with fixed data
                with st.spinner("Updating EFA on fixed dataset…"):
                    S.efa_result  = run_efa(S.df_working, S.efa_result["n_factors"], rotation_method)
                    S.diagnostics = diagnose_loadings(S.efa_result["loadings"],
                                                      S.efa_result["communalities"],
                                                      loading_threshold, communality_threshold)
                    S.efa_done = True
                st.success("✓ Working dataset updated with fixes. Proceed to Step 5.")
                st.rerun()

    # ── Show fix log ─────────────────────────────────────────────────
    if S.fix_log is not None:
        fix_df = S.fix_log

        n_fixed   = int((fix_df["Status"] == "✔ Fixed").sum())
        n_clean   = int((fix_df["Status"] == "✓ Clean").sum())
        n_total   = len(fix_df)

        mc1, mc2, mc3 = st.columns(3)
        mc1.markdown(f'<div class="metric-card"><div class="metric-val" style="color:#34d399;">{n_fixed}</div><div class="metric-label">Variables Fixed</div></div>', unsafe_allow_html=True)
        mc2.markdown(f'<div class="metric-card"><div class="metric-val" style="color:#6c8dfa;">{n_clean}</div><div class="metric-label">Already Clean</div></div>', unsafe_allow_html=True)
        mc3.markdown(f'<div class="metric-card"><div class="metric-val">{n_total}</div><div class="metric-label">Total Variables Retained</div></div>', unsafe_allow_html=True)

        # ── Iteration progress ───────────────────────────────────────
        if S.autofix_efa_result is not None and "iter_log" in S.autofix_efa_result:
            iter_df = S.autofix_efa_result["iter_log"]
            if len(iter_df) > 0:
                st.markdown("#### 🔁 Iterative Fix Progress")
                def _iter_log_html(df):
                    rows = ""
                    for _, r in df.iterrows():
                        flagged = int(r["FlaggedVars"])
                        flag_color = "color:#34d399;" if flagged == 0 else ("color:#fbbf24;" if flagged <= 2 else "color:#f87171;")
                        rows += (
                            f"<tr>"
                            f"<td style='padding:6px 10px;border-bottom:1px solid #1e2540;color:#e2e8f0;text-align:center;'>Pass {int(r['Iteration'])}</td>"
                            f"<td style='padding:6px 10px;border-bottom:1px solid #1e2540;{flag_color}font-weight:700;text-align:center;'>{flagged}</td>"
                            f"<td style='padding:6px 10px;border-bottom:1px solid #1e2540;color:#6c8dfa;text-align:center;'>{r['AvgCommunality']}</td>"
                            f"<td style='padding:6px 10px;border-bottom:1px solid #1e2540;color:#a78bfa;font-size:.82rem;'>{r['FixedThisPass']}</td>"
                            f"</tr>"
                        )
                    return f"""
<div style="overflow-x:auto;">
<table style="width:100%;border-collapse:collapse;font-size:.83rem;font-family:Inter,sans-serif;">
<thead><tr style="background:#1e2540;">
  <th style="padding:8px 10px;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;text-align:center;">Pass</th>
  <th style="padding:8px 10px;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;text-align:center;">Still Flagged</th>
  <th style="padding:8px 10px;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;text-align:center;">Avg Communality</th>
  <th style="padding:8px 10px;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Variables Worked On</th>
</tr></thead>
<tbody>{rows}</tbody>
</table></div>"""
                st.markdown(_iter_log_html(iter_df), unsafe_allow_html=True)

        st.markdown("#### 🔧 Fix Log — Variable-by-Variable")

        def _fix_log_html(df):
            rows = ""
            for _, r in df.iterrows():
                status = r["Status"]
                if "Fixed" in status:
                    s_color = "color:#34d399;font-weight:700;"
                    s_icon  = "✔"
                else:
                    s_color = "color:#64748b;"
                    s_icon  = "✓"
                final_issue = str(r.get("FinalIssue", "—"))
                issue_color = "color:#34d399;" if final_issue in ("OK", "—") else "color:#fbbf24;"
                rows += (
                    f"<tr>"
                    f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;color:#e2e8f0;font-weight:600;'>{r['Variable']}</td>"
                    f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;color:#f87171;font-size:.82rem;'>{r['OriginalDataIssues']}</td>"
                    f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;color:#a78bfa;font-size:.82rem;'>{r['FixesApplied']}</td>"
                    f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;{issue_color}font-size:.82rem;'>{final_issue}</td>"
                    f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;color:#6c8dfa;text-align:center;'>{r['FinalCommunality']}</td>"
                    f"<td style='padding:7px 10px;border-bottom:1px solid #1e2540;{s_color}'>{s_icon} {status.replace('✔ ','').replace('✓ ','')}</td>"
                    f"</tr>"
                )
            return f"""
<div style="overflow-x:auto;">
<table style="width:100%;border-collapse:collapse;font-size:.83rem;font-family:Inter,sans-serif;">
<thead>
<tr style="background:#1e2540;">
  <th style="padding:8px 10px;text-align:left;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Variable</th>
  <th style="padding:8px 10px;text-align:left;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Raw Data Issues</th>
  <th style="padding:8px 10px;text-align:left;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Fixes Applied</th>
  <th style="padding:8px 10px;text-align:left;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Final EFA Issue</th>
  <th style="padding:8px 10px;text-align:center;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Final Communality</th>
  <th style="padding:8px 10px;text-align:left;color:#6c8dfa;font-size:.75rem;text-transform:uppercase;letter-spacing:1px;">Status</th>
</tr>
</thead>
<tbody>
{rows}
</tbody>
</table>
</div>"""
        st.markdown(_fix_log_html(fix_df), unsafe_allow_html=True)

        # ── Before / After EFA comparison ───────────────────────────
        if S.autofix_efa_result is not None:
            st.markdown("#### 📊 EFA Comparison: Before vs After Fix")
            efa_b = S.efa_result
            efa_a = S.autofix_efa_result["efa"]
            diag_b = S.diagnostics
            diag_a = S.autofix_efa_result["diag"]

            n_issues_before = int(diag_b["RecommendDrop"].sum())
            n_issues_after  = int(diag_a["RecommendDrop"].sum())
            comm_before     = round(float(efa_b["communalities"].mean()), 3)
            comm_after      = round(float(efa_a["communalities"].mean()), 3)
            comm_delta      = round(comm_after - comm_before, 3)
            comm_delta_str  = f"+{comm_delta}" if comm_delta >= 0 else str(comm_delta)
            comm_delta_color = "#34d399" if comm_delta >= 0 else "#f87171"

            cc1, cc2, cc3, cc4 = st.columns(4)
            cc1.markdown(f'<div class="metric-card"><div class="metric-val" style="color:#f87171;">{n_issues_before}</div><div class="metric-label">Flagged (Before)</div></div>', unsafe_allow_html=True)
            cc2.markdown(f'<div class="metric-card"><div class="metric-val" style="color:#34d399;">{n_issues_after}</div><div class="metric-label">Flagged (After)</div></div>', unsafe_allow_html=True)
            cc3.markdown(f'<div class="metric-card"><div class="metric-val" style="color:#6c8dfa;">{comm_after}</div><div class="metric-label">Avg Communality (After)</div></div>', unsafe_allow_html=True)
            cc4.markdown(f'<div class="metric-card"><div class="metric-val" style="color:{comm_delta_color};">{comm_delta_str}</div><div class="metric-label">Communality Δ</div></div>', unsafe_allow_html=True)

            t_before, t_after = st.tabs(["Communalities: Before", "Communalities: After"])
            with t_before:
                st.plotly_chart(plot_communalities(efa_b["communalities"], communality_threshold), use_container_width=True, key="comm_before_fix")
            with t_after:
                st.plotly_chart(plot_communalities(efa_a["communalities"], communality_threshold), use_container_width=True, key="comm_after_fix")

        # ── Fixed dataset preview & download ───────────────────────
        if S.df_autofix is not None:
            st.markdown("#### 📋 Fixed Dataset Preview")
            st.dataframe(S.df_autofix.head(10).round(4), use_container_width=True)
            st.markdown(f'<div class="info-box">✅ <b>{len(S.df_autofix.columns)} variables × {len(S.df_autofix)} rows</b> — all variables retained, problematic ones have been transformed in-place.</div>', unsafe_allow_html=True)

            # Always-available free download of fixed dataset (no credit gate here — it's part of the analysis flow)
            fixed_csv = S.df_autofix.to_csv(index=False).encode("utf-8")
            st.download_button(
                label="⬇ Download Fixed Dataset (CSV)",
                data=fixed_csv,
                file_name="efactor_autofix_dataset.csv",
                mime="text/csv",
                use_container_width=True,
                key="dl_fixed_inline",
            )

    # ── Traditional manual drop section ────────────────────────────
    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)
    st.markdown("**— or — manually drop variables (traditional approach)**", unsafe_allow_html=True)

    available = S.df_working.columns.tolist()
    to_drop = st.multiselect("Select variables to drop (optional)", available,
                              default=[v for v in problem_vars if v in available])
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        if st.button("▶ Apply Drops & Re-run EFA", use_container_width=True):
            if to_drop:
                S.df_working   = S.df_working.drop(columns=to_drop)
                S.dropped_vars = list(set(S.dropped_vars + to_drop))
            with st.spinner("Re-running EFA…"):
                S.efa_result  = run_efa(S.df_working, n_factors, rotation_method)
                S.diagnostics = diagnose_loadings(S.efa_result["loadings"],
                                                  S.efa_result["communalities"],
                                                  loading_threshold, communality_threshold)
            st.success(f"✓ EFA re-run. Working set: {len(S.df_working.columns)} variables.")
    with col_d2:
        if st.button("↩ Restore All Variables", use_container_width=True):
            S.df_working   = S.df_original.copy()
            S.dropped_vars = []
            S.efa_result   = None
            S.diagnostics  = None
            S.efa_done     = False
            S.df_autofix   = None
            S.fix_log      = None
            S.autofix_efa_result = None
            st.rerun()
    if S.dropped_vars:
        st.markdown(f'<div class="info-box">📋 Dropped so far: <b>{", ".join(S.dropped_vars)}</b></div>', unsafe_allow_html=True)

st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# STEP 5 — CFA
# ══════════════════════════════════════════════════════════════════
st.markdown('<div class="step-badge"><span class="step-num">5</span> Confirmatory Factor Analysis (CFA)</div>', unsafe_allow_html=True)

if S.efa_result is None:
    st.markdown('<div class="warn-box">Complete EFA (Step 3) first.</div>', unsafe_allow_html=True)
else:
    model_str, factor_vars = build_cfa_model(S.efa_result["loadings"], loading_threshold)
    with st.expander("📝 Auto-generated CFA Model", expanded=False):
        edited_model = st.text_area("Edit model specification if needed:", value=model_str, height=160)
    if st.button("▶ Run CFA", use_container_width=True):
        if not edited_model.strip():
            st.error("Model specification is empty.")
        else:
            with st.spinner("Fitting CFA model…"):
                S.cfa_result = run_cfa(S.df_working, edited_model)
                if S.cfa_result["success"]:
                    S.fit_assessment = assess_cfa_fit(S.cfa_result["fit_indices"], cfa_thresholds)
                    S.cfa_result["model_str"] = edited_model
    if S.cfa_result:
        if not S.cfa_result["success"]:
            st.error(f"CFA failed: {S.cfa_result['error']}")
        else:
            fa = S.fit_assessment
            pill2 = "pill-pass" if fa["overall_pass"] else "pill-fail"
            text2 = "✓ MODEL FIT ADEQUATE" if fa["overall_pass"] else "✗ MODEL FIT INADEQUATE"
            st.markdown(f'<b>Fit:</b> <span class="{pill2}">{text2}</span> — {fa["n_pass"]}/{fa["n_total"]} indices passed', unsafe_allow_html=True)
            st.plotly_chart(plot_fit_indices(fa), use_container_width=True)
            fit_records = [dict(Index=idx, Value=d["value"],
                                Threshold=f"{d['direction']} {d['threshold']}",
                                Status="✓ PASS" if d["pass_"] else "✗ FAIL")
                           for idx, d in fa["indices"].items()]
            st.dataframe(pd.DataFrame(fit_records), use_container_width=True, hide_index=True)
            with st.expander("📊 Parameter Estimates"):
                if S.cfa_result["estimates"] is not None:
                    st.dataframe(S.cfa_result["estimates"], use_container_width=True)
            if not fa["overall_pass"]:
                st.markdown("#### 🔧 Modification Suggestions")
                for s in get_modification_suggestions(fa):
                    st.markdown(f'<div class="warn-box">⚠️ {s}</div>', unsafe_allow_html=True)
                st.markdown('<div class="info-box">💡 Return to Step 4 to drop additional items, or Step 3 to re-extract with different n_factors, then re-run CFA.</div>', unsafe_allow_html=True)

st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# STEP 6 — SYNTHETIC DATA
# ══════════════════════════════════════════════════════════════════
st.markdown('<div class="step-badge"><span class="step-num">6</span> Synthetic Data Generation</div>', unsafe_allow_html=True)

if S.efa_result is None:
    st.markdown('<div class="warn-box">Complete EFA (Step 3) first.</div>', unsafe_allow_html=True)
else:
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("#### Factor-Structure Preserving")
        st.markdown('<div class="info-box">Simulates latent factor scores × loadings + unique variance. Preserves <b>psychometric structure</b>. Recommended for structural validity studies.</div>', unsafe_allow_html=True)
        if st.button("▶ Generate (Factor-Based)", use_container_width=True):
            with st.spinner("Simulating from factor structure…"):
                S.synthetic_factor = generate_factor_based(S.df_working, S.efa_result, n_samples=syn_n, seed=syn_seed)
                S.syn_validation = validate_synthetic(S.df_working, S.synthetic_factor)
            st.success(f"✓ {syn_n} synthetic observations generated.")
    with c2:
        st.markdown("#### Correlation Preserving")
        st.markdown('<div class="info-box">Multivariate normal from empirical <b>covariance matrix</b>. Faster, preserves pairwise correlations but not latent structure explicitly.</div>', unsafe_allow_html=True)
        if st.button("▶ Generate (Correlation-Based)", use_container_width=True):
            with st.spinner("Sampling from multivariate normal…"):
                S.synthetic_corr = generate_correlation_based(S.df_working, n_samples=syn_n, seed=syn_seed)
                S.syn_validation = validate_synthetic(S.df_working, S.synthetic_corr)
            st.success(f"✓ {syn_n} synthetic observations generated.")

    syn_display = S.synthetic_factor if S.synthetic_factor is not None else S.synthetic_corr
    if syn_display is not None:
        t1, t2, t3 = st.tabs(["Preview","Validation Summary","Distribution Comparison"])
        with t1: st.dataframe(syn_display.head(10).round(3), use_container_width=True)
        with t2:
            if S.syn_validation is not None:
                st.dataframe(S.syn_validation, use_container_width=True, hide_index=True)
        with t3: st.plotly_chart(plot_synthetic_comparison(S.df_working, syn_display), use_container_width=True)

st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# STEP 7 — EXPORT  (credit-gated)
# ══════════════════════════════════════════════════════════════════
st.markdown('<div class="step-badge"><span class="step-num">7</span> Export Results</div>', unsafe_allow_html=True)

if S.efa_result is None:
    st.markdown('<div class="warn-box">Complete EFA (Step 3) to enable exports.</div>', unsafe_allow_html=True)
else:
    syn_export = S.synthetic_factor if S.synthetic_factor is not None else S.synthetic_corr
    syn_label  = "factor_based" if S.synthetic_factor is not None else "correlation_based"

    # ── Trial: show lock ─────────────────────────────────────────
    if is_trial:
        st.markdown("""
        <div class="lock-box">
          <div style="font-family:Inter,sans-serif;font-size:15px;font-weight:700;
                      color:#6c8dfa;margin-bottom:8px;">◈ Export is a Paid Feature</div>
          <div style="font-family:Inter,sans-serif;font-size:13px;color:#94a3b8;line-height:1.7;margin-bottom:12px;">
            You're on the Free Trial — all analysis above is fully functional and unlimited.<br>
            To download datasets and reports, purchase an access key with credits.<br><br>
            <b style="color:#e2e8f0;">Credit costs:</b>&nbsp;
            ≤ 300 rows = 1 credit &nbsp;·&nbsp;
            301–1,000 rows = 2 credits &nbsp;·&nbsp;
            &gt; 1,000 rows = 5 credits &nbsp;·&nbsp;
            Word report (.docx) = 1 credit
          </div>
        </div>
        """, unsafe_allow_html=True)
        _, upg_col, _ = st.columns([1, 2, 1])
        with upg_col:
            st.link_button("◈ Get Access Key →", "https://your-payment-link", use_container_width=True)
        st.stop()

    # ── Paid: show exports ───────────────────────────────────────
    current_credits = st.session_state.get("credits", 0)
    REPORT_COST = 1

    col_e1, col_e2, col_e3 = st.columns(3)

    with col_e1:
        st.markdown("##### 🗃️ Cleaned Dataset")
        n_clean = len(S.df_working)
        cost_clean = export_credit_cost(n_clean)
        st.markdown(f'<div class="info-box" style="font-size:.82rem;">{n_clean} rows × {len(S.df_working.columns)} cols<br><b>Cost: {cost_clean} credit{"s" if cost_clean>1 else ""}</b></div>', unsafe_allow_html=True)
        if current_credits < cost_clean:
            st.markdown(f'<div class="warn-box" style="font-size:.82rem;">⚠️ Need {cost_clean} credits ({current_credits} remaining).</div>', unsafe_allow_html=True)
            st.link_button("Top up →", "https://your-payment-link", use_container_width=True)
        else:
            if st.button(f"⬇ Download Cleaned Data ({cost_clean} cr)", use_container_width=True, key="dl_clean_btn"):
                new_bal = deduct_credits(access_key, cost_clean)
                st.session_state["credits"] = new_bal
                st.download_button("⬇ cleaned_data.csv",
                                   data=S.df_working.to_csv(index=False).encode("utf-8"),
                                   file_name="efactor_cleaned_data.csv", mime="text/csv",
                                   use_container_width=True, key="dl_clean_actual")
                st.success(f"✓ {cost_clean} credit(s) deducted. Balance: {new_bal}")

    with col_e2:
        st.markdown("##### 🧪 Synthetic Dataset")
        if syn_export is not None:
            n_syn = len(syn_export)
            cost_syn = export_credit_cost(n_syn)
            st.markdown(f'<div class="info-box" style="font-size:.82rem;">{n_syn} rows × {len(syn_export.columns)} cols<br><b>Cost: {cost_syn} credit{"s" if cost_syn>1 else ""}</b></div>', unsafe_allow_html=True)
            if current_credits < cost_syn:
                st.markdown(f'<div class="warn-box" style="font-size:.82rem;">⚠️ Need {cost_syn} credits ({current_credits} remaining).</div>', unsafe_allow_html=True)
                st.link_button("Top up →", "https://your-payment-link", use_container_width=True)
            else:
                if st.button(f"⬇ Download Synthetic Data ({cost_syn} cr)", use_container_width=True, key="dl_syn_btn"):
                    new_bal = deduct_credits(access_key, cost_syn)
                    st.session_state["credits"] = new_bal
                    st.download_button(f"⬇ synthetic_{syn_label}.csv",
                                       data=syn_export.to_csv(index=False).encode("utf-8"),
                                       file_name=f"efactor_synthetic_{syn_label}.csv", mime="text/csv",
                                       use_container_width=True, key="dl_syn_actual")
                    st.success(f"✓ {cost_syn} credit(s) deducted. Balance: {new_bal}")
        else:
            st.markdown('<div class="warn-box">Generate synthetic data in Step 6 first.</div>', unsafe_allow_html=True)

    with col_e3:
        st.markdown("##### 📄 Analysis Report")
        has_cfa = S.cfa_result is not None and S.cfa_result["success"]
        if has_cfa:
            st.markdown(f'<div class="info-box" style="font-size:.82rem;">Full EFA + CFA Word report (.docx)<br><b>Cost: {REPORT_COST} credit</b></div>', unsafe_allow_html=True)
            if current_credits < REPORT_COST:
                st.markdown(f'<div class="warn-box" style="font-size:.82rem;">⚠️ Insufficient credits.</div>', unsafe_allow_html=True)
                st.link_button("Top up →", "https://your-payment-link", use_container_width=True)
            else:
                if st.button(f"🔨 Build Word Report ({REPORT_COST} cr)", use_container_width=True, key="build_rpt"):
                    with st.spinner("Compiling Word report…"):
                        S.report_docx = generate_docx_report(
                            original_df=S.df_original, cleaned_df=S.df_working,
                            suitability=S.suitability, efa_result=S.efa_result,
                            diagnostics=S.diagnostics, dropped_vars=S.dropped_vars,
                            cfa_result=S.cfa_result, fit_assessment=S.fit_assessment,
                            cfa_thresholds=cfa_thresholds, synthetic_validation=S.syn_validation,
                            model_str=S.cfa_result.get("model_str",""))
                    new_bal = deduct_credits(access_key, REPORT_COST)
                    st.session_state["credits"] = new_bal
                    st.success(f"✓ Report ready. {REPORT_COST} credit deducted. Balance: {new_bal}")
                if S.report_docx:
                    st.download_button("⬇ efa_cfa_report.docx",
                                       data=S.report_docx,
                                       file_name="efactor_report.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       use_container_width=True)
        else:
            st.markdown('<div class="warn-box">Run CFA (Step 5) to enable report generation.</div>', unsafe_allow_html=True)

    # ── ZIP bundle ────────────────────────────────────────────────
    st.divider()
    st.markdown("##### 📦 Full Export Bundle (.zip)")
    zip_rows = max(len(S.df_working), len(syn_export) if syn_export is not None else 0)
    zip_cost = export_credit_cost(zip_rows) + (REPORT_COST if S.report_docx else 0)
    st.markdown(f'<div class="info-box" style="font-size:.82rem;">Cleaned data + synthetic data + Word report + loadings CSV + model spec<br><b>Cost: {zip_cost} credits</b></div>', unsafe_allow_html=True)
    if current_credits < zip_cost:
        st.markdown(f'<div class="warn-box" style="font-size:.82rem;">⚠️ Need {zip_cost} credits ({current_credits} remaining).</div>', unsafe_allow_html=True)
        st.link_button("Top up credits →", "https://your-payment-link", use_container_width=True)
    else:
        if st.button(f"⬇ Build & Download ZIP ({zip_cost} cr)", use_container_width=True, key="dl_zip_btn"):
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                zf.writestr("cleaned_data.csv", S.df_working.to_csv(index=False))
                if syn_export is not None:
                    zf.writestr(f"synthetic_{syn_label}.csv", syn_export.to_csv(index=False))
                if S.report_docx:
                    zf.writestr("efa_cfa_report.docx", S.report_docx)
                if S.cfa_result and S.cfa_result.get("model_str"):
                    zf.writestr("cfa_model.txt", S.cfa_result["model_str"])
                if S.efa_result:
                    zf.writestr("efa_loadings.csv", S.efa_result["loadings"].round(4).to_csv())
                    zf.writestr("efa_communalities.csv", S.efa_result["communalities"].round(4).to_frame().to_csv())
            zip_buffer.seek(0)
            new_bal = deduct_credits(access_key, zip_cost)
            st.session_state["credits"] = new_bal
            ts_str = datetime.now().strftime("%Y%m%d_%H%M")
            st.download_button("⬇ efactor_bundle.zip",
                               data=zip_buffer.getvalue(),
                               file_name=f"efactor_bundle_{ts_str}.zip",
                               mime="application/zip", key="dl_zip_actual")
            st.success(f"✓ {zip_cost} credits deducted. Balance: {new_bal}")


# ══════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════
st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)
st.markdown("""
<div style="text-align:center;padding:8px 0 24px;">
  <div style="font-family:Inter,sans-serif;font-size:14px;font-weight:700;
               background:linear-gradient(90deg,#6c8dfa,#a78bfa);
               -webkit-background-clip:text;-webkit-text-fill-color:transparent;
               margin-bottom:8px;">◈ EFActor</div>
  <div style="font-family:Inter,sans-serif;font-size:11px;color:#374151;">
    Psychometric Analysis Platform &nbsp;·&nbsp; Built for researchers.
    Interpret results within your theoretical framework and research design.
  </div>
  <div style="font-family:Inter,sans-serif;font-size:10px;color:#1e2540;margin-top:8px;">
    © 2025 EFActor
  </div>
</div>
""", unsafe_allow_html=True)
